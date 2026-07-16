from datetime import date, datetime, time
from io import BytesIO
from unittest.mock import patch

from django.contrib.auth import get_user_model
from django.contrib.auth.models import Permission
from django.core.exceptions import ValidationError
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import SimpleTestCase, TestCase
from django.urls import reverse
from django.utils import timezone
from docx import Document
from PIL import Image

from .forms import DrappImportForm, validate_clinical_upload
from .models import Attachment, AttachmentKind, CoverageType, Encounter, EncounterStatus, GeneratedReport, Patient, ReferringPhysician, ReportType, SpirometryResult, StudyType, VitalSigns, WalkTest
from .pdf_intake import (
    build_analysis_from_text,
    build_spirometry_analysis,
    snapshot_with_computed_age,
    extract_patient_snapshot_from_text,
    extract_spirometry_numbers_from_text,
)
from .services import build_reports_for_encounter, construir_informe_espirometria, normalizar_medico
from .views import (
    extract_drapp_rows_from_browser_ocr,
    extract_drapp_rows_from_ocr_lines,
    extract_drapp_rows_from_text,
    format_physician_display_name,
    infer_coverage_type,
    import_drapp_rows,
    get_patient_age_value,
    save_generated_report_artifacts,
    sort_dashboard_encounters,
    unique_encounters_by_patient_day,
)


def grant_clinic_permissions(user, *codenames):
    user.user_permissions.add(*Permission.objects.filter(codename__in=codenames, content_type__app_label="clinic"))


def make_png_upload(name="capture.png"):
    stream = BytesIO()
    Image.new("RGB", (8, 8), color=(245, 250, 252)).save(stream, format="PNG")
    return SimpleUploadedFile(name, stream.getvalue(), content_type="image/png")


class ClinicalUploadValidationTests(SimpleTestCase):
    def test_rejects_pdf_extension_with_non_pdf_content(self):
        uploaded = SimpleUploadedFile("resultado.pdf", b"not a pdf", content_type="application/pdf")

        with self.assertRaisesMessage(ValidationError, "no es un PDF valido"):
            validate_clinical_upload(uploaded)

    def test_rejects_clinical_file_over_fifteen_megabytes(self):
        uploaded = SimpleUploadedFile(
            "resultado.pdf",
            b"%PDF-1.4\n" + (b"0" * (15 * 1024 * 1024)),
            content_type="application/pdf",
        )

        with self.assertRaisesMessage(ValidationError, "supera el limite de 15 MB"):
            validate_clinical_upload(uploaded)

    def test_drapp_screenshot_is_verified_as_a_real_image(self):
        invalid_form = DrappImportForm(files={"screenshot": SimpleUploadedFile("drapp.png", b"fake")})
        valid_form = DrappImportForm(files={"screenshot": make_png_upload()})

        self.assertFalse(invalid_form.is_valid())
        self.assertIn("imagen esta danada", invalid_form.errors["screenshot"][0])
        self.assertTrue(valid_form.is_valid(), valid_form.errors)

    def test_rejects_image_with_excessive_pixel_dimensions(self):
        uploaded = SimpleUploadedFile("drapp.png", b"fake", content_type="image/png")
        with patch("clinic.forms.Image.open") as image_open:
            image_open.return_value.size = (50_000, 50_000)
            with self.assertRaisesMessage(ValidationError, "dimensiones demasiado grandes"):
                validate_clinical_upload(uploaded)


class PhysicianNameFormattingTests(SimpleTestCase):
    def test_keeps_female_doctor_prefix_when_formatting_for_storage(self):
        self.assertEqual(format_physician_display_name("Dra. Maria Perez"), "DRA. Maria Perez")
        self.assertEqual(format_physician_display_name("DR.A Maria Perez"), "DRA. Maria Perez")

    def test_keeps_female_doctor_prefix_when_printing_reports(self):
        self.assertEqual(normalizar_medico("Dra. Maria Perez"), "DRA. MARIA PEREZ")
        self.assertEqual(normalizar_medico("DR.A Maria Perez"), "DRA. MARIA PEREZ")
        self.assertEqual(normalizar_medico("Dr. Juan Perez"), "DR. JUAN PEREZ")


class PatientAgeSourceTests(SimpleTestCase):
    def test_birth_date_takes_priority_over_inconsistent_reported_age(self):
        patient = Patient(birth_date=date(1970, 11, 23), age_reported=23)

        with patch("clinic.views.timezone.localdate", return_value=date(2026, 7, 13)):
            self.assertEqual(get_patient_age_value(patient), 55)


class DashboardOrderingTests(SimpleTestCase):
    def test_attended_order_precedes_pending_schedule_order(self):
        attended_first = Encounter(
            attended=True,
            attended_at=timezone.make_aware(datetime(2026, 7, 13, 15, 5)),
            encounter_time=time(15, 20),
        )
        attended_second = Encounter(
            attended=True,
            attended_at=timezone.make_aware(datetime(2026, 7, 13, 15, 10)),
            encounter_time=time(15, 0),
        )
        pending_early = Encounter(attended=False, encounter_time=time(14, 30))
        pending_late = Encounter(attended=False, encounter_time=time(16, 0))

        ordered = sort_dashboard_encounters(
            [pending_late, attended_second, pending_early, attended_first]
        )

        self.assertEqual(
            ordered,
            [attended_first, attended_second, pending_early, pending_late],
        )


class SpirometryReportTextTests(SimpleTestCase):
    def test_obstructive_pattern_uses_small_airways_text(self):
        text = construir_informe_espirometria("Obstructivo", "leve", "")

        self.assertIn("obstrucción leve", text)
        self.assertIn("pequeñas vías respiratorias aéreas", text)

    def test_mixed_pattern_uses_split_wording(self):
        text = construir_informe_espirometria("Mixto", "severa", "moderadamente severa")

        self.assertIn("patrón mixto", text)
        self.assertIn(" Restricción Moderadamente severa.", text)
        self.assertIn(" Obstrucción Severa a las pequeñas vías respiratorias aéreas.", text)
        self.assertNotIn("\n\n", text)

    def test_restrictive_pattern_keeps_general_airways_text(self):
        text = construir_informe_espirometria("Restrictivo", "", "moderada")

        self.assertIn("respiratorias", text.lower())
        self.assertNotIn("peque", text.lower())


class BronchodilatorAnalysisTests(SimpleTestCase):
    def test_uses_ers_ats_change_over_ten_percent_of_predicted(self):
        analysis = build_spirometry_analysis(
            {
                "fvc": {"lln": 0.4, "predicted": 1.0, "best": 0.5, "percent": 50, "post": 0.61, "post_percent": 61},
                "fev1": {"lln": 0.4, "predicted": 1.0, "best": 0.5, "percent": 50, "post": 0.55, "post_percent": 55},
                "fev1_fvc": {"lln": 70.0, "predicted": 80.0, "best": 100.0, "percent": 125},
            }
        )

        self.assertTrue(analysis["bronchodilator_positive"])
        self.assertIn("criterio ERS/ATS 2022", analysis["bronchodilator_reason"])
        self.assertIn("valor predicho", analysis["bronchodilator_reason"])

    def test_detects_positive_bronchodilator_response_from_post_values(self):
        analysis = build_spirometry_analysis(
            {
                "fvc": {"lln": 2.0, "predicted": 3.0, "best": 2.1, "percent": 70, "post": 2.45, "post_percent": 82, "change_percent": 17},
                "fev1": {"lln": 1.8, "predicted": 2.5, "best": 1.9, "percent": 76, "post": 2.16, "post_percent": 86, "change_percent": 14},
                "fev1_fvc": {"lln": 70.0, "predicted": 80.0, "best": 90.0, "percent": 112, "post": 92.0, "post_percent": 115, "change_percent": 2},
            }
        )

        self.assertTrue(analysis["bronchodilator_positive"])
        self.assertIn("Broncodilatador positivo", analysis["bronchodilator_reason"])

    def test_does_not_mark_positive_without_required_change(self):
        analysis = build_spirometry_analysis(
            {
                "fvc": {"lln": 2.0, "predicted": 3.0, "best": 2.1, "percent": 70, "post": 2.2, "post_percent": 73, "change_percent": 5},
                "fev1": {"lln": 1.8, "predicted": 2.5, "best": 1.9, "percent": 76, "post": 2.02, "post_percent": 80, "change_percent": 6},
                "fev1_fvc": {"lln": 70.0, "predicted": 80.0, "best": 90.0, "percent": 112, "post": 91.0, "post_percent": 114, "change_percent": 1},
            }
        )

        self.assertFalse(analysis["bronchodilator_positive"])
        self.assertEqual(analysis["bronchodilator_reason"], "")


class CoverageInferenceTests(SimpleTestCase):
    def test_only_particular_stays_particular(self):
        self.assertEqual(infer_coverage_type("Particular"), "Particular")

    def test_any_other_coverage_becomes_mutual(self):
        self.assertEqual(infer_coverage_type("grassi"), "Mutual")
        self.assertEqual(infer_coverage_type("DOSEP"), "Mutual")
        self.assertEqual(infer_coverage_type(""), "Mutual")


class DrappImportParsingTests(SimpleTestCase):
    def test_extracts_patient_when_status_lines_come_before_name(self):
        rows = extract_drapp_rows_from_ocr_lines(
            [
                {"text": "Jueves 4 - junio 2026", "y": 0},
                {"text": "15:00", "y": 40},
                {"text": "hace 2 dias", "y": 70},
                {"text": "Reservado", "y": 100},
                {"text": "Oruetta , Ramona", "y": 130},
                {"text": "+54 2657 61 3457", "y": 160},
                {"text": "Particular", "y": 190},
                {"text": "Link de Pago", "y": 220},
                {"text": "Centro Respiratorio Integral", "y": 250},
            ]
        )

        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]["patient_name"], "ORUETTA, RAMONA")
        self.assertEqual(rows[0]["coverage_raw"], "Particular")
        self.assertEqual(rows[0]["agenda_date"], date(2026, 6, 4))

    def test_keeps_importing_when_study_column_is_not_visible(self):
        rows = extract_drapp_rows_from_ocr_lines(
            [
                {"text": "Jueves 4 - junio 2026", "y": 0},
                {"text": "16:15", "y": 40},
                {"text": "Reservado", "y": 70},
                {"text": "Chavez9/144, Marcelo", "y": 100},
                {"text": "+542664648159", "y": 130},
                {"text": "DOSEP", "y": 160},
                {"text": "29.936.370", "y": 190},
            ]
        )

        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]["patient_name"], "CHAVEZ, MARCELO")
        self.assertEqual(rows[0]["coverage_raw"], "DOSEP")
        self.assertEqual(rows[0]["practice_raw"], "")
        self.assertEqual(rows[0]["dni"], "29936370")

    def test_extracts_all_rows_from_notebook_screenshot_ocr(self):
        rows = extract_drapp_rows_from_ocr_lines(
            [
                {"text": "Viernes 5 - junio 2026 - Todos 5 Reservados 5 En Espera 0 En consulta 0 Atendidos 0 Ausentes 0 Cancelados 0 Pendiente 0", "y": 0},
                {"text": "@ 09:20 hace 2 dias ?+542657555332 Avila3/915, Maria Del Carmen 16.133.118 Particular Link de Pago Espirometria Espirometria Espirometria, Piguillem Gustavo Centro Respiratorio Integral", "y": 40},
                {"text": "Reservado", "y": 64},
                {"text": "@ 09:40 hace 15 horas +542658 40 7539 PERALTA, MARCELA 22.822.428 Particular Link de Pago Cicloespirometria Espirometria Espirometria, Piguillem Gustavo Centro Respiratorio Integral", "y": 98},
                {"text": "Reservado", "y": 124},
                {"text": "\u246010:00 Fredes, Josefa PAMI 6/535 Cicloespirometria Espirometria", "y": 160},
                {"text": "Reservado hace 6 dias +542657610914 5.182.672 Link de Pago Espirometria, Piguillem Gustavo Centro Respiratorio Integral", "y": 186},
                {"text": "\u2460 11:00 hace 3 dias Bigner, Maria Estela +5426575809643.697.049 DOSEP Link de Pago Cicloespirometria Espirometria Espirometria, Piguillem Gustavo Centro Respiratorio Integral", "y": 220},
                {"text": "Reservado", "y": 246},
                {"text": "\u246011:15 hace 1 dia +542657 65 4467 PAREDES, CARLOS FENIX 10.705.164 Particular Link de Pago Cicloespirometria Espirometria Espirometria, Piguillem Gustavo Centro Respiratorio Integral", "y": 282},
                {"text": "Reservado", "y": 308},
            ]
        )

        self.assertEqual(len(rows), 5)
        self.assertEqual(
            [(row["datetime_raw"], row["patient_name"]) for row in rows],
            [
                ("09:20", "AVILA, MARIA DEL CARMEN"),
                ("09:40", "PERALTA, MARCELA"),
                ("10:00", "FREDES, JOSEFA"),
                ("11:00", "BIGNER, MARIA ESTELA"),
                ("11:15", "PAREDES, CARLOS FENIX"),
            ],
        )
        self.assertEqual(rows[2]["coverage_raw"], "PAMI")
        self.assertEqual(rows[2]["dni"], "5182672")
        self.assertEqual(rows[3]["coverage_raw"], "DOSEP")
        self.assertEqual(rows[3]["dni"], "3697049")
        self.assertEqual(rows[4]["dni"], "10705164")
        self.assertEqual(rows[0]["agenda_date"], date(2026, 6, 5))

    def test_ocr_like_raw_text_falls_back_to_capture_parser(self):
        rows = extract_drapp_rows_from_text(
            "\n".join(
                [
                    "Viernes 5 - junio 2026",
                    "@ 09:20 hace 2 dias +542657555332 Avila3/915, Maria Del Carmen 16.133.118 Particular",
                    "@ 09:40 hace 15 horas +542658 40 7539 PERALTA, MARCELA 22.822.428 Particular",
                    "\u246010:00 Fredes, Josefa PAMI 6/535 Cicloespirometria",
                    "Reservado hace 6 dias +542657610914 5.182.672",
                    "\u2460 11:00 hace 3 dias Bigner, Maria Estela +5426575809643.697.049 DOSEP",
                    "\u246011:15 hace 1 dia +542657 65 4467 PAREDES, CARLOS FENIX 10.705.164 Particular",
                ]
            )
        )

        self.assertEqual(len(rows), 5)
        self.assertEqual(rows[2]["patient_name"], "FREDES, JOSEFA")
        self.assertEqual(rows[3]["dni"], "3697049")

    def test_trims_study_fragments_and_false_prefixes_from_patient_name(self):
        rows = extract_drapp_rows_from_ocr_lines(
            [
                {"text": "Jueves 11 - junio 2026", "y": 0},
                {"text": "15:55 ORTIZ, MARTA LILIANA 16.484.284 Link de Pago Cicloespirometria Espirometria Espirometria, Piguillem Gustavo Centro Respiratorio Integral", "y": 40},
                {"text": "Reservado hace 6 horas Particular", "y": 64},
                {"text": "16:30 SUAREZ, RODRIGO ESPI 37.797.996 Particular", "y": 98},
                {"text": "16:55 OR DEVIA, CARLOS EVA 13.435.809 PAMI ESPIROMETRIA", "y": 130},
                {"text": "17:15 ESPIROMETRIA, DIM CE 6.484.284 Particular", "y": 160},
                {"text": "18:00 Mercado, Rosario 54.002.437 DOSEP 8/11 Cicloespirometria Espirometria", "y": 190},
            ]
        )

        self.assertEqual(
            [(row["datetime_raw"], row["patient_name"]) for row in rows],
            [
                ("15:55", "ORTIZ, MARTA LILIANA"),
                ("16:30", "SUAREZ, RODRIGO"),
                ("16:55", "DEVIA, CARLOS EVA"),
                ("18:00", "MERCADO, ROSARIO"),
            ],
        )
        self.assertEqual(rows[0]["dni"], "16484284")
        self.assertEqual(rows[1]["dni"], "37797996")
        self.assertEqual(rows[2]["dni"], "13435809")

    def test_browser_ocr_structured_columns_keep_name_dni_and_coverage_clean(self):
        payload = [
            {
                "text": "15:55 ORTIZ, MARTA LILIANA Link de Pago Cicloespirometria",
                "y": 40,
                "items": [
                    {"text": "15:55", "x": 60, "y": 40},
                    {"text": "ORTIZ,", "x": 210, "y": 40},
                    {"text": "MARTA", "x": 290, "y": 40},
                    {"text": "LILIANA", "x": 370, "y": 40},
                    {"text": "Link", "x": 700, "y": 40},
                    {"text": "de", "x": 745, "y": 40},
                    {"text": "Pago", "x": 780, "y": 40},
                    {"text": "Cicloespirometria", "x": 980, "y": 40},
                ],
            },
            {
                "text": "16.484.284",
                "y": 68,
                "items": [
                    {"text": "16.484.284", "x": 420, "y": 68},
                ],
            },
            {
                "text": "16:30 Suarez , Rodrigo +54 2657 29 0348 37.797.996 Particular",
                "y": 110,
                "items": [
                    {"text": "16:30", "x": 60, "y": 110},
                    {"text": "Suarez", "x": 210, "y": 110},
                    {"text": ",", "x": 300, "y": 110},
                    {"text": "Rodrigo", "x": 325, "y": 110},
                    {"text": "+54", "x": 215, "y": 138},
                    {"text": "2657", "x": 255, "y": 138},
                    {"text": "29", "x": 320, "y": 138},
                    {"text": "0348", "x": 355, "y": 138},
                    {"text": "37.797.996", "x": 430, "y": 138},
                    {"text": "Particular", "x": 710, "y": 110},
                ],
            },
            {
                "text": "16:55 OR DEVIA, CARLOS EVARISTO PAMI",
                "y": 180,
                "items": [
                    {"text": "16:55", "x": 60, "y": 180},
                    {"text": "OR", "x": 190, "y": 180},
                    {"text": "DEVIA,", "x": 220, "y": 180},
                    {"text": "CARLOS", "x": 320, "y": 180},
                    {"text": "EVARISTO", "x": 405, "y": 180},
                    {"text": "PAMI", "x": 710, "y": 180},
                    {"text": "+54", "x": 215, "y": 208},
                    {"text": "2657", "x": 255, "y": 208},
                    {"text": "44", "x": 320, "y": 208},
                    {"text": "8093", "x": 355, "y": 208},
                    {"text": "13.435.809", "x": 430, "y": 208},
                ],
            },
        ]

        rows = extract_drapp_rows_from_browser_ocr(__import__("json").dumps(payload))

        self.assertEqual(
            [(row["datetime_raw"], row["patient_name"], row["dni"], row["coverage_raw"]) for row in rows],
            [
                ("15:55", "ORTIZ, MARTA LILIANA", "16484284", ""),
                ("16:30", "SUAREZ, RODRIGO", "37797996", "Particular"),
                ("16:55", "DEVIA, CARLOS EVARISTO", "13435809", "PAMI"),
            ],
        )


class DrappImportViewTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(username="drapp-test", password="secret123")
        grant_clinic_permissions(self.user, "manage_agenda")

    def test_falls_back_to_uploaded_screenshot_when_browser_ocr_returns_no_rows(self):
        self.client.force_login(self.user)
        screenshot = make_png_upload("drapp.png")
        screenshot_rows = [
            {
                "patient_name": "PERALTA, MARCELA",
                "coverage_raw": "Particular",
                "practice_raw": "Cicloespirometria",
                "datetime_raw": "09:40",
                "phone": "+542658407539",
                "dni": "22822428",
                "agenda_date": date(2026, 6, 5),
            }
        ]

        with (
            patch("clinic.views.extract_drapp_rows_from_text", return_value=[]),
            patch("clinic.views.extract_drapp_rows_from_browser_ocr", return_value=[]),
            patch("clinic.views.extract_drapp_rows_from_screenshot", return_value=screenshot_rows) as screenshot_mock,
        ):
            response = self.client.post(
                reverse("clinic:dashboard"),
                {
                    "action": "import_drapp",
                    "raw_text": "texto pobre del OCR del navegador",
                    "ocr_lines_json": '[{"text":"ruido","y":0}]',
                    "screenshot": screenshot,
                },
            )

        self.assertEqual(response.status_code, 200)
        screenshot_mock.assert_called_once()
        self.assertContains(response, "Revisar antes de agregar")
        self.assertContains(response, "PERALTA, MARCELA")
        self.assertEqual(Encounter.objects.count(), 0)

    def test_prefers_browser_ocr_rows_over_plain_raw_text_from_same_capture(self):
        self.client.force_login(self.user)
        browser_rows = [
            {
                "patient_name": "ORTIZ, MARTA LILIANA",
                "coverage_raw": "Particular",
                "practice_raw": "Cicloespirometria",
                "datetime_raw": "15:55",
                "phone": "",
                "dni": "16484284",
                "agenda_date": date(2026, 6, 11),
            },
            {
                "patient_name": "SUAREZ, RODRIGO",
                "coverage_raw": "Particular",
                "practice_raw": "Cicloespirometria",
                "datetime_raw": "16:30",
                "phone": "+542657292348",
                "dni": "37797896",
                "agenda_date": date(2026, 6, 11),
            },
        ]

        with (
            patch(
                "clinic.views.extract_drapp_rows_from_text",
                return_value=[
                    {
                        "patient_name": "ESPIROMETRIA, DINOS CENTRO I RESPIRATORIO CL INTEGRAL I",
                        "coverage_raw": "Particular",
                        "practice_raw": "Cicloespirometria",
                        "datetime_raw": "15:55",
                        "phone": "",
                        "dni": "",
                        "agenda_date": date(2026, 6, 11),
                    }
                ],
            ) as text_mock,
            patch("clinic.views.extract_drapp_rows_from_browser_ocr", return_value=browser_rows),
        ):
            response = self.client.post(
                reverse("clinic:dashboard"),
                {
                    "action": "import_drapp",
                    "raw_text": "15:55 ORTIZ, MARTA LILIANA Particular Espirometria Centro Respiratorio Integral",
                    "ocr_lines_json": '[{"text":"15:55 ORTIZ, MARTA LILIANA","y":40}]',
                },
            )

        self.assertEqual(response.status_code, 200)
        text_mock.assert_not_called()
        self.assertContains(response, "ORTIZ, MARTA LILIANA")
        self.assertContains(response, "SUAREZ, RODRIGO")
        self.assertNotContains(response, "ESPIROMETRIA, DINOS")
        self.assertEqual(Encounter.objects.count(), 0)

    def test_confirm_preview_is_the_only_step_that_creates_encounters(self):
        self.client.force_login(self.user)
        browser_rows = [
            {
                "patient_name": "ORTIZ, MARTA LILIANA",
                "coverage_raw": "PAMI",
                "practice_raw": "Cicloespirometria",
                "datetime_raw": "15:55",
                "phone": "+542657123456",
                "dni": "16484284",
                "agenda_date": date(2026, 6, 11),
            }
        ]
        with patch("clinic.views.extract_drapp_rows_from_browser_ocr", return_value=browser_rows):
            preview_response = self.client.post(
                reverse("clinic:dashboard"),
                {"action": "import_drapp", "ocr_lines_json": '[{"text":"fila","y":1}]'},
            )

        self.assertEqual(preview_response.status_code, 200)
        self.assertEqual(Encounter.objects.count(), 0)
        token = preview_response.context["import_preview_token"]
        confirm_response = self.client.post(
            reverse("clinic:dashboard"),
            {
                "action": "confirm_drapp_import",
                "import_preview_token": token,
                "row_0_selected": "1",
                "row_0_patient_name": "ORTIZ, MARTA LILIANA",
                "row_0_dni": "16.484.284",
                "row_0_coverage": "Mutual",
                "row_0_study": "Ciclometria",
                "row_0_date": "2026-06-12",
                "row_0_time": "15:55",
            },
        )

        self.assertRedirects(confirm_response, reverse("clinic:dashboard"))
        encounter = Encounter.objects.get()
        self.assertEqual(encounter.patient.dni, "16484284")
        self.assertEqual(encounter.coverage_type, CoverageType.MUTUAL)
        self.assertEqual(encounter.encounter_date, date(2026, 6, 12))
        self.assertEqual(encounter.encounter_time, time(15, 55))

    def test_confirm_preview_rejects_missing_date_instead_of_using_today(self):
        self.client.force_login(self.user)
        browser_rows = [
            {
                "patient_name": "ORTIZ, MARTA LILIANA",
                "coverage_raw": "Particular",
                "practice_raw": "Espirometria",
                "datetime_raw": "15:55",
                "phone": "",
                "dni": "16484284",
                "agenda_date": None,
            }
        ]
        with patch("clinic.views.extract_drapp_rows_from_browser_ocr", return_value=browser_rows):
            preview_response = self.client.post(
                reverse("clinic:dashboard"),
                {"action": "import_drapp", "ocr_lines_json": '[{"text":"fila","y":1}]'},
            )

        token = preview_response.context["import_preview_token"]
        response = self.client.post(
            reverse("clinic:dashboard"),
            {
                "action": "confirm_drapp_import",
                "import_preview_token": token,
                "row_0_selected": "1",
                "row_0_patient_name": "ORTIZ, MARTA LILIANA",
                "row_0_dni": "16484284",
                "row_0_coverage": "Particular",
                "row_0_study": "Espirometria",
                "row_0_time": "15:55",
            },
        )

        self.assertRedirects(response, reverse("clinic:dashboard"))
        self.assertEqual(Encounter.objects.count(), 0)


class DashboardInlineUpdateTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(username="agenda-test", password="secret123")
        grant_clinic_permissions(self.user, "manage_agenda")
        self.client.force_login(self.user)
        self.patient = Patient.objects.create(full_name="SUAREZ, MARIA JESUS", dni="10731742")
        self.encounter = Encounter.objects.create(
            patient=self.patient,
            encounter_date=date(2026, 6, 5),
            study_type=StudyType.CICLOMETRIA,
            coverage_type=CoverageType.PARTICULAR,
            status=EncounterStatus.PENDIENTE,
            created_by=self.user,
            updated_by=self.user,
        )

    def test_inline_time_update_persists_and_returns_normalized_value(self):
        response = self.client.post(
            reverse("clinic:dashboard"),
            {
                "action": "inline_update",
                "encounter_id": self.encounter.pk,
                "field_name": "encounter_time",
                "value": "15:30",
            },
            HTTP_X_REQUESTED_WITH="XMLHttpRequest",
        )

        self.assertEqual(response.status_code, 200)
        self.encounter.refresh_from_db()
        self.assertEqual(self.encounter.encounter_time, time(15, 30))
        payload = response.json()
        self.assertEqual(payload["value"], "15:30")
        self.assertEqual(payload["encounter_time"], "15:30")

    def test_dni_owned_by_another_patient_never_reassigns_encounter(self):
        other_patient = Patient.objects.create(full_name="OTRO, PACIENTE", dni="99888777")

        response = self.client.post(
            reverse("clinic:dashboard"),
            {
                "action": "update_dni",
                "encounter_id": self.encounter.pk,
                "patient_dni": "99.888.777",
            },
            HTTP_X_REQUESTED_WITH="XMLHttpRequest",
        )

        self.assertEqual(response.status_code, 409)
        self.encounter.refresh_from_db()
        self.patient.refresh_from_db()
        self.assertEqual(self.encounter.patient_id, self.patient.pk)
        self.assertEqual(self.patient.dni, "10731742")
        self.assertEqual(other_patient.encounters.count(), 0)

    def test_inline_time_update_accepts_compact_manual_value(self):
        response = self.client.post(
            reverse("clinic:dashboard"),
            {
                "action": "inline_update",
                "encounter_id": self.encounter.pk,
                "field_name": "encounter_time",
                "value": "1530",
            },
            HTTP_X_REQUESTED_WITH="XMLHttpRequest",
        )

        self.assertEqual(response.status_code, 200)
        self.encounter.refresh_from_db()
        self.assertEqual(self.encounter.encounter_time, time(15, 30))
        self.assertEqual(response.json()["value"], "15:30")

    def test_rest_vitals_are_saved_together_and_mark_encounter_attended(self):
        self.encounter.no_show = True
        self.encounter.status = EncounterStatus.NO_LLEGO
        self.encounter.save(update_fields=["no_show", "status", "updated_at"])

        response = self.client.post(
            reverse("clinic:dashboard"),
            {
                "action": "save_vitals_group",
                "encounter_id": self.encounter.pk,
                "vitals_group": "rest",
                "so2": "100",
                "fc": "77",
            },
            HTTP_X_REQUESTED_WITH="XMLHttpRequest",
        )

        self.assertEqual(response.status_code, 200)
        self.encounter.refresh_from_db()
        self.assertTrue(self.encounter.attended)
        self.assertFalse(self.encounter.no_show)
        self.assertEqual(self.encounter.status, EncounterStatus.CARGADA)
        self.assertEqual(self.encounter.vital_signs.so2_rest, 100)
        self.assertEqual(self.encounter.vital_signs.fc_rest, 77)

    def test_invalid_vitals_batch_rolls_back_both_values(self):
        self.encounter.no_show = True
        self.encounter.status = EncounterStatus.NO_LLEGO
        self.encounter.save(update_fields=["no_show", "status", "updated_at"])
        VitalSigns.objects.create(encounter=self.encounter, so2_rest=95, fc_rest=70)

        response = self.client.post(
            reverse("clinic:dashboard"),
            {
                "action": "save_vitals_group",
                "encounter_id": self.encounter.pk,
                "vitals_group": "rest",
                "so2": "101",
                "fc": "80",
            },
            HTTP_X_REQUESTED_WITH="XMLHttpRequest",
        )

        self.assertEqual(response.status_code, 400)
        self.encounter.refresh_from_db()
        self.encounter.vital_signs.refresh_from_db()
        self.assertEqual(self.encounter.vital_signs.so2_rest, 95)
        self.assertEqual(self.encounter.vital_signs.fc_rest, 70)
        self.assertFalse(self.encounter.attended)
        self.assertTrue(self.encounter.no_show)

    def test_individual_vital_update_is_rejected(self):
        response = self.client.post(
            reverse("clinic:dashboard"),
            {
                "action": "inline_update",
                "encounter_id": self.encounter.pk,
                "field_name": "so2_rest",
                "value": "95",
            },
            HTTP_X_REQUESTED_WITH="XMLHttpRequest",
        )

        self.assertEqual(response.status_code, 400)
        self.assertFalse(VitalSigns.objects.filter(encounter=self.encounter, so2_rest__isnull=False).exists())

    def test_inline_referring_physician_can_create_new_doctor_name(self):
        response = self.client.post(
            reverse("clinic:dashboard"),
            {
                "action": "inline_update",
                "encounter_id": self.encounter.pk,
                "field_name": "referring_physician",
                "value": "Dr. Pepito Perez",
            },
            HTTP_X_REQUESTED_WITH="XMLHttpRequest",
        )

        self.assertEqual(response.status_code, 200)
        self.encounter.refresh_from_db()
        self.assertIsNotNone(self.encounter.referring_physician)
        self.assertEqual(self.encounter.referring_physician.full_name, "DR. Pepito Perez")
        self.assertTrue(ReferringPhysician.objects.filter(full_name__iexact="DR. Pepito Perez").exists())

    def test_dashboard_uses_explicit_save_for_physician_search(self):
        with patch("clinic.views.timezone.localdate", return_value=date(2026, 6, 5)):
            response = self.client.get(reverse("clinic:dashboard"))

        self.assertEqual(response.status_code, 200)
        html = response.content.decode()
        self.assertIn('data-physician-form', html)
        self.assertIn('data-physician-suggestions', html)
        self.assertIn('Escribi y elegi un doctor de la lista.', html)
        self.assertNotIn('physician-save-button', html)
        self.assertNotIn('data-inline-submit data-physician-search', html)

    def test_dashboard_renders_manual_save_buttons_for_vitals(self):
        with patch("clinic.views.timezone.localdate", return_value=date(2026, 6, 5)):
            response = self.client.get(reverse("clinic:dashboard"))

        self.assertEqual(response.status_code, 200)
        html = response.content.decode()
        self.assertIn('name="action" value="save_vitals_group"', html)
        self.assertIn('data-vitals-batch="1" data-vitals-group="rest"', html)
        self.assertIn('data-vitals-batch="1" data-vitals-group="post"', html)
        self.assertIn('name="so2"', html)
        self.assertIn('name="fc"', html)
        self.assertIn('title="Guardar SO2 y FC en reposo"', html)
        self.assertIn('title="Guardar SO2 y FC post"', html)

    def test_dashboard_ignores_missing_storage_files_in_latest_report(self):
        result = SpirometryResult.objects.create(
            encounter=self.encounter,
            respiratory_pattern="Normal",
        )
        VitalSigns.objects.create(encounter=self.encounter, so2_rest=95, fc_rest=66)
        artifacts = build_reports_for_encounter(self.encounter)
        save_generated_report_artifacts(self.encounter, artifacts, self.user)
        report_attachment = self.encounter.generated_reports.first().attachment
        report_attachment.file.storage.delete(report_attachment.file.name)

        with patch("clinic.views.timezone.localdate", return_value=date(2026, 6, 5)):
            response = self.client.get(reverse("clinic:dashboard"))

        self.assertEqual(response.status_code, 200)
        html = response.content.decode()
        self.assertIn("SUAREZ, MARIA JESUS", html)


class DashboardQuickAddTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(username="quick-add-test", password="secret123")
        grant_clinic_permissions(self.user, "manage_agenda")
        self.client.force_login(self.user)

    def test_quick_add_without_time_uses_current_local_time(self):
        with (
            patch("clinic.views.timezone.localdate", return_value=date(2026, 6, 18)),
            patch("clinic.views.timezone.localtime", return_value=datetime(2026, 6, 18, 14, 37, 44)),
        ):
            response = self.client.post(
                reverse("clinic:dashboard"),
                {
                    "patient_name": "Paciente Hora Auto",
                    "patient_dni": "",
                    "encounter_time": "",
                    "study_type": StudyType.CICLOMETRIA,
                    "coverage_type": CoverageType.PARTICULAR,
                    "distance_meters": "200",
                    "borg_final": "0",
                    "completed": "on",
                },
            )

        self.assertRedirects(response, reverse("clinic:dashboard"))
        encounter = Encounter.objects.get(patient__full_name="PACIENTE HORA AUTO")
        self.assertEqual(encounter.encounter_date, date(2026, 6, 18))
        self.assertEqual(encounter.encounter_time, time(14, 37))
        self.assertFalse(encounter.attended)
        self.assertTrue(encounter.no_show)
        self.assertEqual(encounter.status, EncounterStatus.NO_LLEGO)

    def test_quick_add_keeps_manual_time_when_present(self):
        with patch("clinic.views.timezone.localtime", return_value=datetime(2026, 6, 18, 14, 37, 44)):
            response = self.client.post(
                reverse("clinic:dashboard"),
                {
                    "patient_name": "Paciente Hora Manual",
                    "encounter_time": "09:15",
                    "study_type": StudyType.CICLOMETRIA,
                    "coverage_type": CoverageType.PARTICULAR,
                    "distance_meters": "200",
                    "borg_final": "0",
                    "completed": "on",
                },
            )

        self.assertRedirects(response, reverse("clinic:dashboard"))
        encounter = Encounter.objects.get(patient__full_name="PACIENTE HORA MANUAL")
        self.assertEqual(encounter.encounter_time, time(9, 15))
        self.assertTrue(encounter.no_show)

    def test_quick_add_with_incomplete_rest_vitals_stays_not_attended(self):
        response = self.client.post(
            reverse("clinic:dashboard"),
            {
                "patient_name": "Paciente Con Signos",
                "encounter_time": "10:00",
                "study_type": StudyType.CICLOMETRIA,
                "coverage_type": CoverageType.PARTICULAR,
                "so2_rest": "95",
                "distance_meters": "200",
                "borg_final": "0",
                "completed": "on",
                "no_show": "on",
            },
        )

        self.assertRedirects(response, reverse("clinic:dashboard"))
        encounter = Encounter.objects.get(patient__full_name="PACIENTE CON SIGNOS")
        self.assertFalse(encounter.attended)
        self.assertTrue(encounter.no_show)
        self.assertEqual(encounter.status, EncounterStatus.NO_LLEGO)

    def test_quick_add_with_complete_rest_vitals_starts_attended(self):
        response = self.client.post(
            reverse("clinic:dashboard"),
            {
                "patient_name": "Paciente Con Signos Completos",
                "encounter_time": "10:10",
                "study_type": StudyType.CICLOMETRIA,
                "coverage_type": CoverageType.PARTICULAR,
                "so2_rest": "95",
                "fc_rest": "74",
                "distance_meters": "200",
                "borg_final": "0",
                "completed": "on",
                "no_show": "on",
            },
        )

        self.assertRedirects(response, reverse("clinic:dashboard"))
        encounter = Encounter.objects.get(patient__full_name="PACIENTE CON SIGNOS COMPLETOS")
        self.assertTrue(encounter.attended)
        self.assertFalse(encounter.no_show)
        self.assertEqual(encounter.status, EncounterStatus.CARGADA)

    def test_quick_add_accepts_custom_referring_physician_name(self):
        response = self.client.post(
            reverse("clinic:dashboard"),
            {
                "patient_name": "Paciente Con Doctor Libre",
                "encounter_time": "11:15",
                "study_type": StudyType.CICLOMETRIA,
                "coverage_type": CoverageType.PARTICULAR,
                "referring_physician": "Dra. Nueva Propuesta",
                "distance_meters": "200",
                "borg_final": "0",
                "completed": "on",
                "no_show": "on",
            },
        )

        self.assertRedirects(response, reverse("clinic:dashboard"))
        encounter = Encounter.objects.select_related("referring_physician").get(patient__full_name="PACIENTE CON DOCTOR LIBRE")
        self.assertIsNotNone(encounter.referring_physician)
        self.assertEqual(encounter.referring_physician.full_name, "DR. Nueva Propuesta")
        self.assertTrue(ReferringPhysician.objects.filter(full_name__iexact="DR. Nueva Propuesta").exists())

    def test_existing_dni_never_overwrites_canonical_patient_name(self):
        patient = Patient.objects.create(full_name="PEREZ, JUAN", dni="30111222")

        response = self.client.post(
            reverse("clinic:dashboard"),
            {
                "patient_name": "NOMBRE OCR EQUIVOCADO",
                "patient_dni": "30.111.222",
                "encounter_time": "12:00",
                "study_type": StudyType.CICLOMETRIA,
                "coverage_type": CoverageType.PARTICULAR,
                "distance_meters": "200",
                "borg_final": "0",
                "completed": "on",
            },
        )

        self.assertRedirects(response, reverse("clinic:dashboard"))
        patient.refresh_from_db()
        self.assertEqual(patient.full_name, "PEREZ, JUAN")
        self.assertTrue(Encounter.objects.filter(patient=patient, encounter_time=time(12, 0)).exists())

    def test_quick_add_reuses_deleted_patient_identity_without_restoring_old_encounters(self):
        patient = Patient.objects.create(full_name="PEREZ, JUAN", dni="30111222")
        old_encounter = Encounter.objects.create(
            patient=patient,
            encounter_date=date(2026, 5, 1),
            study_type=StudyType.ESPIROMETRIA,
            coverage_type=CoverageType.PARTICULAR,
        )
        patient.soft_delete(deleted_by=self.user)

        response = self.client.post(
            reverse("clinic:dashboard"),
            {
                "patient_name": "PEREZ, JUAN",
                "patient_dni": "30.111.222",
                "encounter_time": "12:30",
                "study_type": StudyType.CICLOMETRIA,
                "coverage_type": CoverageType.PARTICULAR,
                "distance_meters": "200",
                "borg_final": "0",
                "completed": "on",
            },
        )

        self.assertRedirects(response, reverse("clinic:dashboard"))
        patient.refresh_from_db()
        old_encounter.refresh_from_db()
        self.assertIsNone(patient.deleted_at)
        self.assertIsNotNone(old_encounter.deleted_at)
        self.assertEqual(Encounter.objects.filter(patient=patient).count(), 1)



class DoctorReviewViewTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(username="review-test", password="secret123")
        grant_clinic_permissions(self.user, "review_medically")
        self.patient = Patient.objects.create(full_name="SIPOLLONI, FELISA", dni=None)
        self.encounter = Encounter.objects.create(
            patient=self.patient,
            encounter_date=date(2026, 6, 4),
            encounter_time=time(16, 35),
            study_type=StudyType.CICLOMETRIA,
            coverage_type=CoverageType.PARTICULAR,
            status=EncounterStatus.PENDIENTE,
            created_by=self.user,
            updated_by=self.user,
        )

    def test_list_defaults_to_buenos_aires_calendar_date(self):
        self.client.force_login(self.user)

        with patch("clinic.views.timezone.localdate", return_value=date(2026, 7, 13)):
            response = self.client.get(reverse("clinic:doctor_review_list"))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context["selected_date"], "2026-07-13")
        self.assertEqual(response.context["today"], date(2026, 7, 13))

    def test_list_deduplicates_same_patient_same_day_and_keeps_latest_card(self):
        self.client.force_login(self.user)
        duplicate_patient = Patient.objects.create(full_name="THEO, CORNEJO", dni="59158072")
        Encounter.objects.create(
            patient=duplicate_patient,
            encounter_date=date(2026, 7, 14),
            encounter_time=time(12, 10),
            study_type=StudyType.CICLOMETRIA,
            coverage_type=CoverageType.PARTICULAR,
            status=EncounterStatus.PENDIENTE,
            created_by=self.user,
            updated_by=self.user,
        )
        latest_encounter = Encounter.objects.create(
            patient=duplicate_patient,
            encounter_date=date(2026, 7, 14),
            encounter_time=time(12, 16),
            study_type=StudyType.CICLOMETRIA,
            coverage_type=CoverageType.PARTICULAR,
            status=EncounterStatus.CARGADA,
            attended=True,
            attended_at=timezone.make_aware(datetime(2026, 7, 14, 12, 20)),
            created_by=self.user,
            updated_by=self.user,
        )
        Attachment.objects.create(
            encounter=latest_encounter,
            file_kind=AttachmentKind.PDF_RESULTADO,
            original_name="theo.pdf",
            file="encounters/theo/theo.pdf",
            mime_type="application/pdf",
            uploaded_by=self.user,
        )

        response = self.client.get(reverse("clinic:doctor_review_list"), {"date": "2026-07-14"})

        self.assertEqual(response.status_code, 200)
        html = response.content.decode()
        self.assertEqual(html.count("THEO, CORNEJO"), 1)
        self.assertIn("Revisar PDF", html)
        self.assertNotIn("Abrir ficha", html)

    def test_upload_suggestion_does_not_become_medical_result(self):
        self.client.force_login(self.user)
        pdf_file = SimpleUploadedFile("felisa.pdf", b"%PDF-1.4 fake", content_type="application/pdf")
        analysis = {
            "source": "browser-pdf-text",
            "code": "RM",
            "probability": 99,
            "summary": "99% probable RM. FVC debajo del LIN/LLN.",
            "bronchodilator_positive": True,
            "bronchodilator_reason": "Respuesta positiva sugerida por FEV1.",
            "values": {
                "fvc": {"lln": 1.75, "predicted": 2.69, "best": 1.65, "percent": 61},
                "fev1": {"lln": 1.36, "predicted": 1.98, "best": 1.58, "percent": 80},
                "fev1_fvc": {"lln": 65.5, "predicted": 77.1, "best": 95.8, "percent": 124},
            },
        }

        with patch("clinic.views.build_analysis_for_uploaded_result", return_value=analysis):
            response = self.client.post(
                reverse("clinic:doctor_review_detail", args=[self.encounter.pk]),
                {
                    "pdf_file": pdf_file,
                    "respiratory_result": "",
                    "analysis_payload_json": "{}",
                },
            )

        self.assertRedirects(response, reverse("clinic:doctor_review_detail", args=[self.encounter.pk]))
        self.encounter.refresh_from_db()
        self.assertTrue(self.encounter.attended)
        self.assertFalse(self.encounter.no_show)
        self.assertEqual(self.encounter.status, EncounterStatus.CARGADA)
        self.assertEqual(self.encounter.spirometry_result.suggested_code, "RM")
        self.assertEqual(self.encounter.spirometry_result.respiratory_pattern, "")
        self.assertTrue(self.encounter.spirometry_result.suggested_bronchodilator_positive)
        self.assertFalse(self.encounter.spirometry_result.bronchodilator_positive)

    def test_upload_autofills_missing_document_profile_and_replaces_random_name(self):
        self.client.force_login(self.user)
        self.patient.full_name = "ASDASDAS"
        self.patient.dni = None
        self.patient.save(update_fields=["full_name", "dni", "updated_at"])
        pdf_file = SimpleUploadedFile("palermo.pdf", b"%PDF-1.4 fake", content_type="application/pdf")
        analysis = {
            "source": "browser-pdf-text",
            "code": "N",
            "probability": 99,
            "summary": "99% probable N.",
            "values": {
                "fvc": {"lln": 1.79, "predicted": 2.50, "best": 2.37, "percent": 95},
                "fev1": {"lln": 1.48, "predicted": 2.11, "best": 2.34, "percent": 111},
                "fev1_fvc": {"lln": 67.9, "predicted": 78.6, "best": 98.7, "percent": 125},
            },
            "snapshot": {
                "patient_code": "12345678",
                "dni": "12345678",
                "last_name": "MARTIN",
                "first_name": "PALERMO",
                "full_name": "MARTIN, PALERMO",
                "birth_date": date(1970, 11, 23),
                "age_reported": 23,
                "gender": "Femenino",
                "height_cm": 154,
                "weight_kg": "100",
                "bmi": "42.17",
                "ethnicity": "Caucasico",
                "smoking_status": "No fumador",
            },
        }

        with patch("clinic.views.build_analysis_for_uploaded_result", return_value=analysis):
            response = self.client.post(
                reverse("clinic:doctor_review_detail", args=[self.encounter.pk]),
                {
                    "pdf_file": pdf_file,
                    "respiratory_result": "",
                    "analysis_payload_json": "{}",
                },
            )

        self.assertRedirects(response, reverse("clinic:doctor_review_detail", args=[self.encounter.pk]))
        self.patient.refresh_from_db()
        self.assertEqual(self.patient.full_name, "MARTIN, PALERMO")
        self.assertEqual(self.patient.dni, "12345678")
        self.assertEqual(self.patient.patient_code, "12345678")
        self.assertEqual(self.patient.last_name, "MARTIN")
        self.assertEqual(self.patient.first_name, "PALERMO")
        self.assertEqual(self.patient.birth_date, date(1970, 11, 23))
        self.assertEqual(self.patient.age_reported, 55)
        self.assertEqual(self.patient.gender, "Femenino")
        self.assertEqual(self.patient.height_cm, 154)
        self.assertEqual(str(self.patient.bmi), "42.17")

    def test_save_review_re_reads_existing_pdf_and_autofills_profile(self):
        self.client.force_login(self.user)
        Attachment.objects.create(
            encounter=self.encounter,
            file_kind=AttachmentKind.PDF_RESULTADO,
            original_name="castro.pdf",
            file="encounters/45/castro.pdf",
            mime_type="application/pdf",
            uploaded_by=self.user,
        )
        analysis = {
            "source": "server-pdf-text",
            "code": "RS",
            "probability": 99,
            "summary": "99% probable RS.",
            "values": {
                "fvc": {"lln": 1.96, "predicted": 2.90, "best": 0.68, "percent": 23},
                "fev1": {"lln": 1.62, "predicted": 2.23, "best": 0.59, "percent": 26},
                "fev1_fvc": {"lln": 66.1, "predicted": 77.7, "best": 86.8, "percent": 112},
            },
            "snapshot": {
                "patient_code": "12231324",
                "dni": "12231324",
                "last_name": "CASTRO",
                "first_name": "ARGENTINA",
                "full_name": "CASTRO, ARGENTINA",
                "birth_date": date(1955, 12, 25),
                "age_reported": 70,
                "gender": "Femenino",
                "height_cm": 165,
                "weight_kg": "63",
                "bmi": "23.14",
                "ethnicity": "Caucásico",
            },
        }

        self.patient.full_name = "CASTRO ARGENTINA"
        self.patient.dni = "12231324"
        self.patient.save(update_fields=["full_name", "dni", "updated_at"])
        with patch("clinic.views.build_analysis_for_uploaded_result", return_value=analysis):
            response = self.client.post(
                reverse("clinic:doctor_review_detail", args=[self.encounter.pk]),
                {
                    "pdf_file": "",
                    "respiratory_result": "RS",
                    "analysis_payload_json": "",
                },
            )

        self.assertRedirects(response, reverse("clinic:doctor_review_detail", args=[self.encounter.pk]))
        self.patient.refresh_from_db()
        self.assertEqual(self.patient.gender, "Femenino")
        self.assertEqual(self.patient.birth_date, date(1955, 12, 25))
        self.assertEqual(self.patient.age_reported, 70)
        self.assertEqual(self.patient.height_cm, 165)
        self.assertEqual(str(self.patient.bmi), "23.14")

    def test_save_review_persists_manual_bronchodilator_flag(self):
        self.client.force_login(self.user)

        response = self.client.post(
            reverse("clinic:doctor_review_detail", args=[self.encounter.pk]),
            {
                "pdf_file": "",
                "respiratory_result": "RM",
                "bronchodilator_positive": "on",
                "analysis_payload_json": "",
            },
        )

        self.assertRedirects(response, reverse("clinic:doctor_review_detail", args=[self.encounter.pk]))
        self.encounter.refresh_from_db()
        self.assertTrue(self.encounter.spirometry_result.bronchodilator_positive)


class DoctorReviewNavigationTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(username="doctor-nav", password="secret123")
        grant_clinic_permissions(self.user, "review_medically")
        self.client.force_login(self.user)
        self.today = date(2026, 6, 28)

        self.current_patient = Patient.objects.create(full_name="PACIENTE ACTUAL", dni="11111111")
        self.next_patient = Patient.objects.create(full_name="PACIENTE SIGUIENTE", dni="22222222")
        self.done_patient = Patient.objects.create(full_name="PACIENTE RESUELTO", dni="33333333")

        self.current_encounter = Encounter.objects.create(
            patient=self.current_patient,
            encounter_date=self.today,
            encounter_time=time(15, 0),
            study_type=StudyType.CICLOMETRIA,
            coverage_type=CoverageType.PARTICULAR,
            status=EncounterStatus.CARGADA,
            attended=True,
            attended_at=timezone.make_aware(datetime(2026, 6, 28, 15, 5)),
            created_by=self.user,
            updated_by=self.user,
        )
        self.next_encounter = Encounter.objects.create(
            patient=self.next_patient,
            encounter_date=self.today,
            encounter_time=time(15, 20),
            study_type=StudyType.CICLOMETRIA,
            coverage_type=CoverageType.PARTICULAR,
            status=EncounterStatus.CARGADA,
            attended=True,
            attended_at=timezone.make_aware(datetime(2026, 6, 28, 15, 12)),
            created_by=self.user,
            updated_by=self.user,
        )
        self.done_encounter = Encounter.objects.create(
            patient=self.done_patient,
            encounter_date=self.today,
            encounter_time=time(15, 40),
            study_type=StudyType.CICLOMETRIA,
            coverage_type=CoverageType.PARTICULAR,
            status=EncounterStatus.REVISADA,
            attended=True,
            attended_at=timezone.make_aware(datetime(2026, 6, 28, 15, 30)),
            created_by=self.user,
            updated_by=self.user,
        )

        Attachment.objects.create(
            encounter=self.current_encounter,
            file_kind=AttachmentKind.PDF_RESULTADO,
            original_name="actual.pdf",
            file="encounters/current/actual.pdf",
            mime_type="application/pdf",
            uploaded_by=self.user,
        )
        Attachment.objects.create(
            encounter=self.next_encounter,
            file_kind=AttachmentKind.PDF_RESULTADO,
            original_name="next.pdf",
            file="encounters/next/next.pdf",
            mime_type="application/pdf",
            uploaded_by=self.user,
        )
        Attachment.objects.create(
            encounter=self.done_encounter,
            file_kind=AttachmentKind.PDF_RESULTADO,
            original_name="done.pdf",
            file="encounters/done/done.pdf",
            mime_type="application/pdf",
            uploaded_by=self.user,
        )

    def test_review_detail_exposes_queue_context(self):
        response = self.client.get(reverse("clinic:doctor_review_detail", args=[self.current_encounter.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context["review_queue"]["pending_total"], 2)
        self.assertEqual(response.context["review_queue"]["remaining_after_current"], 2)
        self.assertIsNone(response.context["previous_review_encounter"])
        self.assertEqual(response.context["next_review_encounter"].pk, self.next_encounter.pk)

    def test_review_queue_state_returns_next_pending_patient(self):
        response = self.client.get(reverse("clinic:doctor_review_queue_state", args=[self.current_encounter.pk]))

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload["ok"])
        self.assertEqual(payload["pending_total"], 2)
        self.assertTrue(payload["current_is_pending"])
        self.assertEqual(payload["remaining_after_current"], 2)
        self.assertIsNone(payload["previous_encounter_id"])
        self.assertEqual(payload["previous_encounter_name"], "")
        self.assertEqual(payload["previous_encounter_label"], "")
        self.assertEqual(payload["next_encounter_id"], self.next_encounter.pk)
        self.assertEqual(payload["next_encounter_name"], "PACIENTE SIGUIENTE")
        self.assertEqual(payload["next_encounter_label"], "15:20 - PACIENTE SIGUIENTE")

    def test_queue_decreases_only_after_saving_a_final_result(self):
        first_state = self.client.get(
            reverse("clinic:doctor_review_queue_state", args=[self.current_encounter.pk])
        ).json()
        self.assertEqual(first_state["pending_total"], 2)

        with patch("clinic.views.build_analysis_for_uploaded_result", return_value={}):
            response = self.client.post(
                reverse("clinic:doctor_review_detail", args=[self.current_encounter.pk]),
                {
                    "pdf_file": "",
                    "respiratory_result": "N",
                    "analysis_payload_json": "",
                },
            )

        self.assertRedirects(
            response,
            reverse("clinic:doctor_review_detail", args=[self.current_encounter.pk]),
        )
        self.current_encounter.refresh_from_db()
        self.assertEqual(self.current_encounter.status, EncounterStatus.REVISADA)

        final_state = self.client.get(
            reverse("clinic:doctor_review_queue_state", args=[self.current_encounter.pk])
        ).json()
        self.assertEqual(final_state["pending_total"], 1)
        self.assertFalse(final_state["current_is_pending"])
        self.assertEqual(final_state["next_encounter_id"], self.next_encounter.pk)


class DrappImportDeduplicationTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(username="dedupe-test", password="secret123")

    def create_encounter(self, full_name="ORUETTA, RAMONA", dni="", encounter_date=date(2026, 6, 4)):
        patient = Patient.objects.create(full_name=full_name, dni=dni or None)
        return Encounter.objects.create(
            patient=patient,
            encounter_date=encounter_date,
            encounter_time=time(15, 0),
            study_type=StudyType.CICLOMETRIA,
            coverage_type=CoverageType.PARTICULAR,
            status=EncounterStatus.PENDIENTE,
            created_by=self.user,
            updated_by=self.user,
        )

    def test_skips_same_name_on_same_day_even_if_time_changes(self):
        self.create_encounter()

        created, skipped = import_drapp_rows(
            [
                {
                    "patient_name": "Oruetta, Ramona",
                    "coverage_raw": "Particular",
                    "practice_raw": "Espirometria",
                    "datetime_raw": "16:30",
                    "phone": "",
                    "dni": "",
                    "agenda_date": date(2026, 6, 4),
                }
            ],
            self.user,
        )

        self.assertEqual(created, 0)
        self.assertEqual(skipped, 1)
        self.assertEqual(Encounter.objects.filter(encounter_date=date(2026, 6, 4)).count(), 1)

    def test_skips_same_dni_on_same_day_even_if_name_has_ocr_prefix(self):
        self.create_encounter(full_name="ORUETTA, RAMONA", dni="42999801")

        created, skipped = import_drapp_rows(
            [
                {
                    "patient_name": "O ORUETTA, RAMONA",
                    "coverage_raw": "Particular",
                    "practice_raw": "Cicloespirometria",
                    "datetime_raw": "17:00",
                    "phone": "",
                    "dni": "42.999.801",
                    "agenda_date": date(2026, 6, 4),
                }
            ],
            self.user,
        )

        self.assertEqual(created, 0)
        self.assertEqual(skipped, 1)
        self.assertEqual(Encounter.objects.filter(patient__dni="42999801").count(), 1)

    def test_allows_same_patient_on_another_day(self):
        self.create_encounter(full_name="ORUETTA, RAMONA", dni="42999801")

        created, skipped = import_drapp_rows(
            [
                {
                    "patient_name": "ORUETTA, RAMONA",
                    "coverage_raw": "Particular",
                    "practice_raw": "Cicloespirometria",
                    "datetime_raw": "15:00",
                    "phone": "",
                    "dni": "42999801",
                    "agenda_date": date(2026, 6, 5),
                }
            ],
            self.user,
        )

        self.assertEqual(created, 1)
        self.assertEqual(skipped, 0)
        self.assertEqual(Encounter.objects.filter(patient__dni="42999801").count(), 2)
        imported = Encounter.objects.get(patient__dni="42999801", encounter_date=date(2026, 6, 5))
        self.assertTrue(imported.no_show)
        self.assertFalse(imported.attended)
        self.assertEqual(imported.status, EncounterStatus.NO_LLEGO)

    def test_import_reuses_deleted_patient_identity_without_restoring_old_encounters(self):
        old_encounter = self.create_encounter(full_name="ORUETTA, RAMONA", dni="42999801")
        patient = old_encounter.patient
        patient.soft_delete(deleted_by=self.user)

        created, skipped = import_drapp_rows(
            [
                {
                    "patient_name": "ORUETTA, RAMONA",
                    "coverage_raw": "Particular",
                    "practice_raw": "Cicloespirometria",
                    "datetime_raw": "15:00",
                    "phone": "",
                    "dni": "42.999.801",
                    "agenda_date": date(2026, 6, 5),
                }
            ],
            self.user,
        )

        patient.refresh_from_db()
        old_encounter.refresh_from_db()
        self.assertEqual((created, skipped), (1, 0))
        self.assertIsNone(patient.deleted_at)
        self.assertIsNotNone(old_encounter.deleted_at)
        self.assertTrue(Encounter.objects.filter(patient=patient, encounter_date=date(2026, 6, 5)).exists())

    def test_skips_same_name_with_different_dni_on_same_day(self):
        self.create_encounter(full_name="PEREZ, MARIA", dni="11111111")

        created, skipped = import_drapp_rows(
            [
                {
                    "patient_name": "PEREZ, MARIA",
                    "coverage_raw": "Particular",
                    "practice_raw": "Espirometria",
                    "datetime_raw": "16:00",
                    "phone": "",
                    "dni": "22222222",
                    "agenda_date": date(2026, 6, 4),
                }
            ],
            self.user,
        )

        self.assertEqual((created, skipped), (0, 1))
        self.assertEqual(Encounter.objects.filter(encounter_date=date(2026, 6, 4)).count(), 1)

    def test_rejects_practice_text_as_patient_name(self):
        created, skipped = import_drapp_rows(
            [
                {
                    "patient_name": "ESPIROMETRIA, CENTRO RESPIRATORIO",
                    "coverage_raw": "Particular",
                    "practice_raw": "Espirometria",
                    "datetime_raw": "16:00",
                    "phone": "",
                    "dni": "6484284",
                    "agenda_date": date(2026, 6, 4),
                }
            ],
            self.user,
        )

        self.assertEqual((created, skipped), (0, 1))
        self.assertFalse(Patient.objects.exists())

    def test_display_uniques_existing_duplicate_encounters(self):
        patient = Patient.objects.create(full_name="O ORUETTA, RAMONA", dni="42999801")
        first = Encounter.objects.create(
            patient=patient,
            encounter_date=date(2026, 6, 4),
            encounter_time=time(15, 0),
            study_type=StudyType.CICLOMETRIA,
            coverage_type=CoverageType.PARTICULAR,
            status=EncounterStatus.PENDIENTE,
        )
        Encounter.objects.create(
            patient=patient,
            encounter_date=date(2026, 6, 4),
            encounter_time=time(16, 0),
            study_type=StudyType.ESPIROMETRIA,
            coverage_type=CoverageType.MUTUAL,
            status=EncounterStatus.PENDIENTE,
        )

        unique = unique_encounters_by_patient_day(
            Encounter.objects.select_related("patient").order_by("encounter_time")
        )

        self.assertEqual(unique, [first])


class SpirometryPdfParsingTests(SimpleTestCase):
    def test_restrictive_pdf_table_is_not_called_normal(self):
        analysis = build_analysis_from_text(
            "\n".join(
                [
                    "FVC FEV1 FEV1%",
                    "Fecha de visita 04/06/2026",
                    "Cod. paciente 5999253",
                    "Apellido SIPOLLONI",
                    "Nom. FELISA",
                    "Genero Femenino",
                    "Altura, cm 164",
                    "Peso, kg 88",
                    "BMI 32,72",
                    "Parametros LLN Teor. Best %Teor. Z-score PRE #1 PRE #2 PRE #3 POST %Teor. %Cam",
                    "FVC L 1,75 2,69 1,65* 61 -1,82 1,65 1,54* 57 -7",
                    "FEV1 L 1,36 1,98 1,58* 80 -1,07 1,58 1,49* 75 -6",
                    "FEV1/FVC % 65,5 77,1 95,8* 124 2,66 95,8 96,8* 126 1",
                ]
            ),
            source="browser-pdf-text",
        )

        self.assertEqual(analysis["values"]["fvc"]["lln"], 1.75)
        self.assertEqual(analysis["values"]["fvc"]["best"], 1.65)
        self.assertEqual(analysis["values"]["fvc"]["percent"], 61)
        self.assertEqual(analysis["values"]["fev1_fvc"]["best"], 95.8)
        self.assertEqual(analysis["code"], "RM")

    def test_uses_distinct_fev1_and_ratio_lines_for_mixed_pattern(self):
        values = extract_spirometry_numbers_from_text(
            "\n".join(
                [
                    "FVC 2.50 3.20 1.80 56",
                    "FEV1 1.80 2.60 1.00 38",
                    "FEV1/FVC 70 81 55",
                ]
            )
        )
        analysis = build_spirometry_analysis(values)

        self.assertEqual(values["fev1_fvc"]["lln"], 70)
        self.assertEqual(values["fev1_fvc"]["best"], 55)
        self.assertEqual(analysis["code"], "RMSOS")

    def test_does_not_call_normal_when_key_values_are_missing(self):
        analysis = build_spirometry_analysis(
            {
                "fvc": {"lln": None, "predicted": None, "best": None, "percent": None},
                "fev1": {"lln": None, "predicted": None, "best": None, "percent": None},
                "fev1_fvc": {"lln": None, "predicted": None, "best": None, "percent": None},
            }
        )

        self.assertEqual(analysis["code"], "")
        self.assertIsNone(analysis["probability"])

    def test_rejects_physically_impossible_repeated_rows(self):
        analysis = build_spirometry_analysis(
            {
                "fvc": {"lln": 4.0, "predicted": None, "best": 16.0, "percent": 77.0},
                "fev1": {"lln": 4.0, "predicted": None, "best": 16.0, "percent": 77.0},
                "fev1_fvc": {"lln": 4.0, "predicted": None, "best": 16.0, "percent": 77.0},
            }
        )

        self.assertEqual(analysis["code"], "")
        self.assertIsNone(analysis["probability"])

    def test_reads_patient_profile_without_colons(self):
        snapshot = extract_patient_snapshot_from_text(
            "\n".join(
                [
                    "Cod. paciente 5999253",
                    "Apellido SIPOLLONI",
                    "Nom. FELISA",
                    "Fecha de nacimien 16/10/1949",
                    "Edad 76",
                    "Genero Femenino",
                    "Altura, cm 164",
                    "Peso, kg 88",
                    "BMI 32,72",
                ]
            )
        )

        self.assertEqual(snapshot["patient_code"], "5999253")
        self.assertEqual(snapshot["full_name"], "SIPOLLONI, FELISA")
        self.assertEqual(snapshot["gender"], "Femenino")
        self.assertEqual(snapshot["height_cm"], 164)

    def test_reads_accented_patient_code_label(self):
        snapshot = extract_patient_snapshot_from_text(
            "\n".join(
                [
                    "Código de paciente: 24.990.727",
                    "Apellido: MUÑOZ",
                    "Nombre: JAVIER ALEJANDRO",
                ]
            )
        )

        self.assertEqual(snapshot["patient_code"], "24990727")
        self.assertEqual(snapshot["dni"], "24990727")
        self.assertEqual(snapshot["full_name"], "MUÑOZ, JAVIER ALEJANDRO")

    def test_reads_patient_profile_when_pdf_joins_left_and_right_columns(self):
        snapshot = extract_patient_snapshot_from_text(
            "\n".join(
                [
                    "Fecha de visita 04/06/2026",
                    "Cod. paciente 5999253 Edad 76",
                    "Apellido SIPOLLONI Genero Femenino",
                    "Nom. FELISA Altura, cm 164",
                    "Fecha de nacimien 16/10/1949 Peso, kg 88",
                    "Grupo etnico Caucasico BMI 32,72",
                    "Fuma Paquete-ano",
                ]
            )
        )

        self.assertEqual(snapshot["patient_code"], "5999253")
        self.assertEqual(snapshot["dni"], "5999253")
        self.assertEqual(snapshot["full_name"], "SIPOLLONI, FELISA")
        self.assertEqual(snapshot["birth_date"], date(1949, 10, 16))
        self.assertEqual(snapshot["age_reported"], 76)
        self.assertEqual(snapshot["gender"], "Femenino")
        self.assertEqual(snapshot["height_cm"], 164)
        self.assertEqual(snapshot["weight_kg"], 88)
        self.assertEqual(str(snapshot["bmi"]), "32.72")
        self.assertEqual(snapshot["ethnicity"], "Caucásico")

    def test_reads_vertical_patient_profile_block_and_computes_age_from_birth_date(self):
        snapshot = extract_patient_snapshot_from_text(
            "\n".join(
                [
                    "Fecha de visita 04/06/2026",
                    "Cod. paciente 12345678 Edad 23",
                    "Apellido",
                    "Nom.",
                    "Fecha de nacimien",
                    "Grupo etnico",
                    "Fuma",
                    "Grupo pacientes",
                    "MARTIN",
                    "PALERMO",
                    "23/11/1970",
                    "Caucasico",
                    "No fumador",
                    "Genero Femenino",
                    "Altura, cm 154",
                    "Peso:, kg 100",
                    "BMI 42,17",
                    "Paquete-ano",
                ]
            )
        )
        snapshot = snapshot_with_computed_age(snapshot, date(2026, 6, 5))

        self.assertEqual(snapshot["patient_code"], "12345678")
        self.assertEqual(snapshot["dni"], "12345678")
        self.assertEqual(snapshot["last_name"], "MARTIN")
        self.assertEqual(snapshot["first_name"], "PALERMO")
        self.assertEqual(snapshot["full_name"], "MARTIN, PALERMO")
        self.assertEqual(snapshot["birth_date"], date(1970, 11, 23))
        self.assertEqual(snapshot["age_reported"], 55)
        self.assertEqual(snapshot["gender"], "Femenino")
        self.assertEqual(snapshot["height_cm"], 154)
        self.assertEqual(snapshot["weight_kg"], 100)
        self.assertEqual(str(snapshot["bmi"]), "42.17")
        self.assertEqual(snapshot["ethnicity"], "Caucásico")
        self.assertEqual(snapshot["smoking_status"], "No fumador")


class PrintReportViewTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(username="print-test", password="secret123")
        self.client.force_login(self.user)
        self.patient = Patient.objects.create(full_name="FONTANARI ALICIA NOEMI", dni="129")
        self.encounter = Encounter.objects.create(
            patient=self.patient,
            encounter_date=date(2026, 6, 4),
            encounter_time=time(15, 0),
            study_type=StudyType.CICLOMETRIA,
            coverage_type=CoverageType.MUTUAL,
            status=EncounterStatus.REVISADA,
            created_by=self.user,
            updated_by=self.user,
        )
        VitalSigns.objects.create(
            encounter=self.encounter,
            so2_rest=88,
            fc_rest=64,
            so2_post=65,
            fc_post=117,
        )
        WalkTest.objects.create(
            encounter=self.encounter,
            distance_meters=100,
            completed=False,
            stopped=False,
            symptoms=False,
            borg_final=1,
        )
        SpirometryResult.objects.create(
            encounter=self.encounter,
            respiratory_pattern="Restrictivo",
            restriction_grade="Moderada",
        )

    def test_mutual_patient_print_includes_mutual_packet(self):
        response = self.client.get(reverse("clinic:encounter_print", args=[self.encounter.pk]))

        self.assertEqual(response.status_code, 200)
        html = response.content.decode()
        self.assertIn("Capacidad Vital Lenta", html)
        self.assertIn("PRUEBA DE LOS 6 Y 12 MINUTOS", html)
        self.assertIn("Resultado Espirometria Computarizada", html)
        self.assertIn("Moderadamente reducida", html)
        self.assertIn("dni-value", html)
        self.assertIn("PRUEBA NO NORMAL", html)
        self.assertIn("SO2: 88%", html)
        self.assertIn("FC: 64", html)
        self.assertNotIn("FC: 64%", html)
        self.assertNotIn("desaturacion al esfuerzo", html.lower())

    def test_daily_print_uses_same_mutual_packet(self):
        with patch("clinic.views.timezone.localdate", return_value=date(2026, 6, 4)):
            response = self.client.get(reverse("clinic:daily_print"))

        self.assertEqual(response.status_code, 200)
        html = response.content.decode()
        self.assertIn("@page { size: Letter; margin: 0; }", html)
        self.assertIn("sheet-pdf", html)
        self.assertIn("Capacidad Vital Lenta", html)
        self.assertIn("Moderadamente reducida", html)
        self.assertIn("PRUEBA NO NORMAL", html)
        self.assertIn("SO2: 88%", html)
        self.assertIn("FC: 64", html)
        self.assertNotIn("FC: 64%", html)

    def test_daily_print_prefers_complete_duplicate_over_incomplete_duplicate(self):
        duplicate = Encounter.objects.create(
            patient=self.patient,
            encounter_date=date(2026, 6, 4),
            encounter_time=time(12, 16),
            study_type=StudyType.CICLOMETRIA,
            coverage_type=CoverageType.PARTICULAR,
            status=EncounterStatus.PENDIENTE,
            created_by=self.user,
            updated_by=self.user,
        )
        self.assertIsNotNone(duplicate.pk)

        with patch("clinic.views.timezone.localdate", return_value=date(2026, 6, 4)):
            response = self.client.get(reverse("clinic:daily_print"))

        self.assertEqual(response.status_code, 200)
        html = response.content.decode()
        self.assertNotIn("Faltan datos antes de imprimir todo el dia", html)
        self.assertIn("FONTANARI ALICIA NOEMI", html)
        self.assertIn("SO2: 88%", html)
        self.assertIn("FC: 64", html)

    def test_generated_docx_prints_fc_without_percent_symbol(self):
        artifacts = build_reports_for_encounter(self.encounter)
        doc = Document(BytesIO(artifacts[0].bytes_content))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

        self.assertIn("SO2: 88%", text)
        self.assertIn("FC: 64", text)
        self.assertNotIn("FC: 64%", text)

    def test_generated_mixed_report_keeps_compact_split_wording(self):
        result = self.encounter.spirometry_result
        result.respiratory_pattern = "Mixto"
        result.restriction_grade = "Moderadamente severa"
        result.obstruction_grade = "Severa"
        result.save(
            update_fields=[
                "respiratory_pattern",
                "restriction_grade",
                "obstruction_grade",
                "updated_at",
            ]
        )

        doc = Document(BytesIO(build_reports_for_encounter(self.encounter)[0].bytes_content))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

        self.assertIn("patrón mixto:\n Restricción Moderadamente severa.", text)
        self.assertIn("\n Obstrucción Severa a las pequeñas vías respiratorias aéreas.", text)

    def test_bronchodilator_line_appears_only_when_final_flag_is_positive(self):
        result = self.encounter.spirometry_result
        negative_doc = Document(BytesIO(build_reports_for_encounter(self.encounter)[0].bytes_content))
        negative_text = "\n".join(paragraph.text for paragraph in negative_doc.paragraphs)
        self.assertNotIn("Broncodilatador Positivo", negative_text)

        result.bronchodilator_positive = True
        result.save(update_fields=["bronchodilator_positive", "updated_at"])
        positive_doc = Document(BytesIO(build_reports_for_encounter(self.encounter)[0].bytes_content))
        positive_text = "\n".join(paragraph.text for paragraph in positive_doc.paragraphs)
        self.assertIn("Broncodilatador Positivo", positive_text)

    def test_spirometry_only_never_adds_walk_page(self):
        self.encounter.study_type = StudyType.ESPIROMETRIA
        self.encounter.save(update_fields=["study_type", "updated_at"])

        artifacts = build_reports_for_encounter(self.encounter)
        doc = Document(BytesIO(artifacts[0].bytes_content))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

        self.assertEqual(len(artifacts), 1)
        self.assertEqual(artifacts[0].report_type, ReportType.ESPIROMETRIA)
        self.assertNotIn("PRUEBA DE LOS 6 Y 12 MINUTOS", text)

    def test_generated_walk_table_never_invents_intermediate_measurements(self):
        artifacts = build_reports_for_encounter(self.encounter)
        doc = Document(BytesIO(artifacts[0].bytes_content))
        walk_table = next(table for table in doc.tables if table.rows[0].cells[0].text == "MINUTOS")

        self.assertEqual(len(walk_table.rows), 8)
        self.assertEqual(walk_table.rows[1].cells[0].text, "0")
        self.assertEqual(walk_table.rows[1].cells[1].text, "88")
        self.assertEqual(walk_table.rows[1].cells[2].text, "64")
        self.assertEqual(walk_table.rows[1].cells[3].text, "0")
        self.assertEqual(walk_table.rows[4].cells[0].text, "3")
        self.assertEqual(walk_table.rows[4].cells[1].text, "76")
        self.assertEqual(walk_table.rows[4].cells[2].text, "90")
        self.assertEqual(walk_table.rows[4].cells[3].text, "0")
        self.assertEqual(walk_table.rows[7].cells[0].text, "6")
        self.assertEqual(walk_table.rows[7].cells[1].text, "65")
        self.assertEqual(walk_table.rows[7].cells[2].text, "117")
        self.assertEqual(walk_table.rows[7].cells[3].text, "1")

    def test_default_zero_borg_prints_as_one_at_minute_six(self):
        walk = self.encounter.walk_test
        walk.borg_final = 0
        walk.save(update_fields=["borg_final", "updated_at"])

        artifacts = build_reports_for_encounter(self.encounter)
        doc = Document(BytesIO(artifacts[0].bytes_content))
        walk_table = next(table for table in doc.tables if table.rows[0].cells[0].text == "MINUTOS")

        self.assertEqual(walk_table.rows[1].cells[3].text, "0")
        self.assertEqual(walk_table.rows[6].cells[3].text, "0")
        self.assertEqual(walk_table.rows[7].cells[3].text, "1")

    def test_regenerated_report_keeps_source_snapshot_hash_and_version_chain(self):
        first_artifacts = build_reports_for_encounter(self.encounter)
        save_generated_report_artifacts(self.encounter, first_artifacts, self.user)
        first_report = GeneratedReport.objects.filter(
            encounter=self.encounter,
            report_type=ReportType.COMPLETO,
        ).latest("created_at")

        vital = self.encounter.vital_signs
        vital.so2_rest = 89
        vital.save(update_fields=["so2_rest", "updated_at"])
        second_artifacts = build_reports_for_encounter(self.encounter)
        save_generated_report_artifacts(self.encounter, second_artifacts, self.user)
        latest_report = GeneratedReport.objects.filter(
            encounter=self.encounter,
            report_type=ReportType.COMPLETO,
        ).latest("created_at")

        self.assertEqual(latest_report.supersedes_id, first_report.pk)
        self.assertEqual(len(latest_report.content_sha256), 64)
        self.assertEqual(latest_report.source_snapshot["patient"]["full_name"], "FONTANARI ALICIA NOEMI")
        self.assertEqual(latest_report.source_snapshot["vital_signs"]["so2_rest"], 89)


class AttachmentUrlResilienceTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(username="storage-test", password="secret123")
        grant_clinic_permissions(self.user, "manage_agenda", "review_medically")
        self.client.force_login(self.user)
        self.patient = Patient.objects.create(full_name="PACIENTE STORAGE", dni="30111222")
        self.encounter = Encounter.objects.create(
            patient=self.patient,
            encounter_date=date(2026, 7, 14),
            encounter_time=time(10, 30),
            study_type=StudyType.ESPIROMETRIA,
            coverage_type=CoverageType.PARTICULAR,
            status=EncounterStatus.REVISADA,
            attended=True,
            created_by=self.user,
            updated_by=self.user,
        )
        VitalSigns.objects.create(encounter=self.encounter, so2_rest=96, fc_rest=72)
        SpirometryResult.objects.create(encounter=self.encounter, respiratory_pattern="Normal")
        self.attachment = Attachment.objects.create(
            encounter=self.encounter,
            file_kind=AttachmentKind.FOTO_RESULTADO,
            original_name="resultado-storage.png",
            file="encounters/storage/resultado-storage.png",
            mime_type="image/png",
            uploaded_by=self.user,
        )
        GeneratedReport.objects.create(
            encounter=self.encounter,
            report_type=ReportType.ESPIROMETRIA,
            attachment=self.attachment,
            generated_by=self.user,
        )

    def storage_url_failure(self):
        return patch.object(
            self.attachment.file.storage,
            "url",
            side_effect=RuntimeError("storage url unavailable"),
        )

    def test_doctor_review_renders_when_attachment_url_is_unavailable(self):
        with self.storage_url_failure():
            response = self.client.get(reverse("clinic:doctor_review_detail", args=[self.encounter.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Archivo original no disponible")

    def test_patient_history_renders_when_attachment_url_is_unavailable(self):
        with self.storage_url_failure():
            response = self.client.get(reverse("clinic:patient_detail", args=[self.patient.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Resultado original no disponible")
        self.assertContains(response, "resultado-storage.png (no disponible)")

    def test_encounter_detail_renders_when_attachment_url_is_unavailable(self):
        with self.storage_url_failure():
            response = self.client.get(reverse("clinic:encounter_detail", args=[self.encounter.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Archivo no disponible")
        self.assertContains(response, "resultado-storage.png (no disponible)")

    def test_print_view_renders_without_broken_original_file_link(self):
        with self.storage_url_failure():
            response = self.client.get(reverse("clinic:encounter_print", args=[self.encounter.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, "Abrir resultado original")
        self.assertContains(response, "PACIENTE STORAGE")


class PatientHistoryActionsTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(username="history-test", password="secret123")
        grant_clinic_permissions(self.user, "manage_agenda", "purge_clinical_data")
        self.client.force_login(self.user)
        self.patient = Patient.objects.create(full_name="TEST, PACIENTE", dni="12345678")
        self.encounter = Encounter.objects.create(
            patient=self.patient,
            encounter_date=date(2026, 6, 11),
            encounter_time=time(16, 30),
            study_type=StudyType.CICLOMETRIA,
            coverage_type=CoverageType.PARTICULAR,
            status=EncounterStatus.PENDIENTE,
            created_by=self.user,
            updated_by=self.user,
        )
        VitalSigns.objects.create(encounter=self.encounter, so2_rest=88, fc_rest=64, so2_post=66, fc_post=117)
        WalkTest.objects.create(encounter=self.encounter, distance_meters=100, completed=False, stopped=False, symptoms=False, borg_final=1)

    def test_patient_detail_shows_walk_assessment_and_edit_action(self):
        response = self.client.get(reverse("clinic:patient_detail", args=[self.patient.pk]))

        self.assertEqual(response.status_code, 200)
        html = response.content.decode()
        self.assertIn("PRUEBA NO NORMAL", html)
        self.assertIn("Editar atencion", html)
        self.assertIn("Generar completo", html)
        self.assertIn("Subir documento", html)

    def test_print_without_final_result_is_blocked_instead_of_defaulting_to_normal(self):
        response = self.client.get(reverse("clinic:encounter_print", args=[self.encounter.pk]))

        self.assertRedirects(response, reverse("clinic:encounter_detail", args=[self.encounter.pk]))

    def test_print_is_allowed_when_result_is_set_even_without_generated_report(self):
        self.encounter.study_type = StudyType.ESPIROMETRIA
        self.encounter.save(update_fields=["study_type", "updated_at"])
        SpirometryResult.objects.create(encounter=self.encounter, respiratory_pattern="Normal")

        response = self.client.get(reverse("clinic:encounter_print", args=[self.encounter.pk]))

        self.assertEqual(response.status_code, 200)
        html = response.content.decode()
        self.assertIn("Resultado Espirometria Computarizada", html)
        self.assertIn("El paciente presenta resultados normales.", html)

    def test_patient_detail_can_upload_document_to_specific_encounter(self):
        upload = SimpleUploadedFile(
            "otro-estudio.pdf",
            b"%PDF-1.4 extra",
            content_type="application/pdf",
        )

        response = self.client.post(
            reverse("clinic:patient_detail", args=[self.patient.pk]),
            {
                "action": "upload_patient_document",
                "encounter": self.encounter.pk,
                "file_kind": AttachmentKind.PDF_RESULTADO,
                "file": upload,
            },
        )

        self.assertRedirects(response, reverse("clinic:patient_detail", args=[self.patient.pk]))
        self.assertTrue(
            Attachment.objects.filter(encounter=self.encounter, original_name="otro-estudio.pdf").exists()
        )

    def test_encounter_edit_can_return_to_patient_history(self):
        response = self.client.post(
            reverse("clinic:encounter_edit", args=[self.encounter.pk]),
            {
                "patient_name": "TEST, PACIENTE",
                "patient_dni": "12345678",
                "encounter_time": "16:30",
                "study_type": "Ciclometria",
                "coverage_type": "Particular",
                "referring_physician": "",
                "so2_rest": 90,
                "fc_rest": 70,
                "so2_post": 85,
                "fc_post": 120,
                "distance_meters": 100,
                "completed": "on",
                "stopped": "",
                "symptoms": "",
                "borg_final": 1,
                "respiratory_result": "N",
                "attended": "",
                "no_show": "",
                "return_to": reverse("clinic:patient_detail", args=[self.patient.pk]),
            },
        )

        self.assertRedirects(response, reverse("clinic:patient_detail", args=[self.patient.pk]))

    def test_encounter_edit_preserves_original_date_and_status(self):
        self.encounter.encounter_date = date(2026, 6, 4)
        self.encounter.status = EncounterStatus.INFORME_GENERADO
        self.encounter.coverage_type = CoverageType.PARTICULAR
        self.encounter.save(update_fields=["encounter_date", "status", "coverage_type", "updated_at"])

        response = self.client.post(
            reverse("clinic:encounter_edit", args=[self.encounter.pk]),
            {
                "patient_name": "TEST, PACIENTE",
                "patient_dni": "12345678",
                "encounter_time": "16:30",
                "study_type": "Ciclometria",
                "coverage_type": "Mutual",
                "referring_physician": "",
                "so2_rest": 90,
                "fc_rest": 70,
                "so2_post": 85,
                "fc_post": 120,
                "distance_meters": 100,
                "completed": "on",
                "stopped": "",
                "symptoms": "",
                "borg_final": 1,
                "respiratory_result": "N",
                "attended": "on",
                "no_show": "",
                "return_to": reverse("clinic:patient_detail", args=[self.patient.pk]),
            },
        )

        self.assertRedirects(response, reverse("clinic:patient_detail", args=[self.patient.pk]))
        self.encounter.refresh_from_db()
        self.assertEqual(self.encounter.encounter_date, date(2026, 6, 4))
        self.assertEqual(self.encounter.status, EncounterStatus.INFORME_GENERADO)
        self.assertEqual(self.encounter.coverage_type, CoverageType.MUTUAL)

    def test_encounter_edit_rejects_dni_owned_by_another_patient_atomically(self):
        other_patient = Patient.objects.create(full_name="OTRO, PACIENTE", dni="87654321")

        response = self.client.post(
            reverse("clinic:encounter_edit", args=[self.encounter.pk]),
            {
                "patient_name": "NOMBRE QUE NO DEBE GUARDARSE",
                "patient_dni": "87.654.321",
                "encounter_time": "18:15",
                "study_type": "Espirometria",
                "coverage_type": "Mutual",
                "referring_physician": "",
                "so2_rest": 90,
                "fc_rest": 70,
                "so2_post": 85,
                "fc_post": 120,
                "distance_meters": 100,
                "completed": "on",
                "borg_final": 1,
                "respiratory_result": "N",
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "ya pertenece a OTRO, PACIENTE")
        self.encounter.refresh_from_db()
        self.patient.refresh_from_db()
        other_patient.refresh_from_db()
        self.assertEqual(self.encounter.patient_id, self.patient.pk)
        self.assertEqual(self.encounter.encounter_time, time(16, 30))
        self.assertEqual(self.encounter.coverage_type, CoverageType.PARTICULAR)
        self.assertEqual(self.patient.full_name, "TEST, PACIENTE")
        self.assertEqual(other_patient.full_name, "OTRO, PACIENTE")

    def test_dashboard_delete_sends_encounter_to_trash_even_with_documents(self):
        Attachment.objects.create(
            encounter=self.encounter,
            file_kind=AttachmentKind.PDF_RESULTADO,
            original_name="resultado.pdf",
            file="encounters/protegido/resultado.pdf",
            mime_type="application/pdf",
            uploaded_by=self.user,
        )

        response = self.client.post(
            reverse("clinic:dashboard"),
            {
                "action": "delete_encounter",
                "encounter_id": self.encounter.pk,
            },
            HTTP_X_REQUESTED_WITH="XMLHttpRequest",
        )

        self.assertEqual(response.status_code, 200)
        self.encounter.refresh_from_db()
        self.assertIsNotNone(self.encounter.deleted_at)
        self.assertFalse(Encounter.objects.filter(pk=self.encounter.pk).exists())
        self.assertTrue(Encounter.all_objects.filter(pk=self.encounter.pk).exists())
        self.assertEqual(response.json()["deleted"], True)

    def test_patient_detail_can_delete_encounter_from_history(self):
        response = self.client.post(
            reverse("clinic:patient_detail", args=[self.patient.pk]),
            {
                "action": "delete_encounter",
                "encounter_id": self.encounter.pk,
            },
        )

        self.assertRedirects(response, reverse("clinic:patient_detail", args=[self.patient.pk]))
        self.encounter.refresh_from_db()
        self.assertIsNotNone(self.encounter.deleted_at)
        self.assertFalse(Encounter.objects.filter(pk=self.encounter.pk).exists())
        self.assertTrue(Encounter.all_objects.filter(pk=self.encounter.pk).exists())

    def test_patient_delete_sends_patient_with_clinical_history_to_trash(self):
        response = self.client.post(reverse("clinic:patient_delete", args=[self.patient.pk]))

        self.assertRedirects(response, reverse("clinic:patient_list"))
        self.patient.refresh_from_db()
        self.encounter.refresh_from_db()
        self.assertIsNotNone(self.patient.deleted_at)
        self.assertIsNotNone(self.encounter.deleted_at)
        self.assertFalse(Patient.objects.filter(pk=self.patient.pk).exists())
        self.assertFalse(Encounter.objects.filter(pk=self.encounter.pk).exists())
        self.assertTrue(Patient.all_objects.filter(pk=self.patient.pk).exists())
        self.assertTrue(Encounter.all_objects.filter(pk=self.encounter.pk).exists())

    def test_recycle_bin_can_restore_patient_and_encounter(self):
        self.patient.soft_delete(deleted_by=self.user)

        response = self.client.post(
            reverse("clinic:recycle_bin"),
            {
                "action": "restore_patient",
                "patient_id": self.patient.pk,
            },
        )

        self.assertRedirects(response, reverse("clinic:recycle_bin"))
        self.patient.refresh_from_db()
        self.encounter.refresh_from_db()
        self.assertIsNone(self.patient.deleted_at)
        self.assertIsNone(self.encounter.deleted_at)
        self.assertTrue(Patient.objects.filter(pk=self.patient.pk).exists())
        self.assertTrue(Encounter.objects.filter(pk=self.encounter.pk).exists())

    def test_restoring_patient_does_not_restore_an_older_independent_deletion(self):
        older_encounter = Encounter.objects.create(
            patient=self.patient,
            encounter_date=date(2026, 5, 1),
            study_type=StudyType.ESPIROMETRIA,
            coverage_type=CoverageType.PARTICULAR,
            created_by=self.user,
            updated_by=self.user,
        )
        older_encounter.soft_delete(deleted_by=self.user)
        older_batch = Encounter.all_objects.get(pk=older_encounter.pk).deletion_batch

        self.patient.soft_delete(deleted_by=self.user)
        Patient.all_objects.get(pk=self.patient.pk).restore()

        restored_current = Encounter.all_objects.get(pk=self.encounter.pk)
        still_deleted_older = Encounter.all_objects.get(pk=older_encounter.pk)
        self.assertIsNone(restored_current.deleted_at)
        self.assertIsNotNone(still_deleted_older.deleted_at)
        self.assertEqual(still_deleted_older.deletion_batch, older_batch)

    def test_failed_report_generation_does_not_change_attendance(self):
        self.encounter.no_show = True
        self.encounter.attended = False
        self.encounter.status = EncounterStatus.NO_LLEGO
        self.encounter.save(update_fields=["no_show", "attended", "status", "updated_at"])

        response = self.client.post(
            reverse("clinic:encounter_generate_report", args=[self.encounter.pk]),
            {"next": reverse("clinic:dashboard")},
        )

        self.assertRedirects(response, reverse("clinic:dashboard"))
        self.encounter.refresh_from_db()
        self.assertFalse(self.encounter.attended)
        self.assertTrue(self.encounter.no_show)
        self.assertEqual(self.encounter.status, EncounterStatus.NO_LLEGO)

    def test_report_inconsistency_warning_does_not_change_attendance(self):
        self.encounter.no_show = True
        self.encounter.attended = False
        self.encounter.status = EncounterStatus.NO_LLEGO
        self.encounter.save(update_fields=["no_show", "attended", "status", "updated_at"])
        vital = self.encounter.vital_signs
        vital.so2_post = 40
        vital.save(update_fields=["so2_post", "updated_at"])
        SpirometryResult.objects.create(encounter=self.encounter, respiratory_pattern="Normal")

        response = self.client.post(
            reverse("clinic:encounter_generate_report", args=[self.encounter.pk]),
            HTTP_X_REQUESTED_WITH="XMLHttpRequest",
        )

        self.assertEqual(response.status_code, 409)
        self.encounter.refresh_from_db()
        self.assertFalse(self.encounter.attended)
        self.assertTrue(self.encounter.no_show)
        self.assertEqual(self.encounter.status, EncounterStatus.NO_LLEGO)

    def test_patient_detail_inconsistency_warning_does_not_generate_or_change_attendance(self):
        self.encounter.no_show = True
        self.encounter.attended = False
        self.encounter.status = EncounterStatus.NO_LLEGO
        self.encounter.save(update_fields=["no_show", "attended", "status", "updated_at"])
        vital = self.encounter.vital_signs
        vital.so2_post = 40
        vital.save(update_fields=["so2_post", "updated_at"])
        SpirometryResult.objects.create(encounter=self.encounter, respiratory_pattern="Normal")

        response = self.client.post(
            reverse("clinic:patient_detail", args=[self.patient.pk]),
            {
                "action": "generate_patient_report",
                "encounter_id": self.encounter.pk,
                "report_scope": "complete",
            },
        )

        self.assertRedirects(response, reverse("clinic:patient_detail", args=[self.patient.pk]))
        self.encounter.refresh_from_db()
        self.assertFalse(self.encounter.attended)
        self.assertTrue(self.encounter.no_show)
        self.assertEqual(self.encounter.status, EncounterStatus.NO_LLEGO)
        self.assertFalse(GeneratedReport.objects.filter(encounter=self.encounter).exists())

    def test_report_build_error_does_not_change_attendance(self):
        self.encounter.no_show = True
        self.encounter.attended = False
        self.encounter.status = EncounterStatus.NO_LLEGO
        self.encounter.save(update_fields=["no_show", "attended", "status", "updated_at"])
        SpirometryResult.objects.create(encounter=self.encounter, respiratory_pattern="Normal")

        with patch("clinic.views.build_reports_for_encounter", side_effect=RuntimeError("fallo controlado")):
            response = self.client.post(
                reverse("clinic:encounter_generate_report", args=[self.encounter.pk]),
                HTTP_X_REQUESTED_WITH="XMLHttpRequest",
            )

        self.assertEqual(response.status_code, 500)
        self.encounter.refresh_from_db()
        self.assertFalse(self.encounter.attended)
        self.assertTrue(self.encounter.no_show)
        self.assertEqual(self.encounter.status, EncounterStatus.NO_LLEGO)

    def test_recycle_bin_can_purge_patient_and_associated_encounter(self):
        self.patient.soft_delete(deleted_by=self.user)

        response = self.client.post(
            reverse("clinic:recycle_bin"),
            {
                "action": "purge_patient",
                "patient_id": self.patient.pk,
            },
        )

        self.assertRedirects(response, reverse("clinic:recycle_bin"))
        self.assertFalse(Patient.all_objects.filter(pk=self.patient.pk).exists())
        self.assertFalse(Encounter.all_objects.filter(pk=self.encounter.pk).exists())

    def test_recycle_bin_cannot_purge_an_active_patient(self):
        response = self.client.post(
            reverse("clinic:recycle_bin"),
            {
                "action": "purge_patient",
                "patient_id": self.patient.pk,
            },
        )

        self.assertEqual(response.status_code, 404)
        self.assertTrue(Patient.objects.filter(pk=self.patient.pk).exists())
        self.assertTrue(Encounter.objects.filter(pk=self.encounter.pk).exists())

    def test_patient_detail_can_generate_mutual_report_only(self):
        self.encounter.coverage_type = CoverageType.MUTUAL
        self.encounter.save(update_fields=["coverage_type", "updated_at"])
        SpirometryResult.objects.create(encounter=self.encounter, respiratory_pattern="Normal")

        response = self.client.post(
            reverse("clinic:patient_detail", args=[self.patient.pk]),
            {
                "action": "generate_patient_report",
                "encounter_id": self.encounter.pk,
                "report_scope": "mutual",
            },
        )

        self.assertRedirects(response, reverse("clinic:patient_detail", args=[self.patient.pk]))
        reports = GeneratedReport.objects.filter(encounter=self.encounter)
        self.assertEqual(reports.count(), 1)
        self.assertEqual(reports.get().report_type, ReportType.MUTUAL)

    def test_patient_detail_complete_report_includes_mutual_when_coverage_is_mutual(self):
        self.encounter.coverage_type = CoverageType.MUTUAL
        self.encounter.save(update_fields=["coverage_type", "updated_at"])
        SpirometryResult.objects.create(encounter=self.encounter, respiratory_pattern="Normal")

        response = self.client.post(
            reverse("clinic:patient_detail", args=[self.patient.pk]),
            {
                "action": "generate_patient_report",
                "encounter_id": self.encounter.pk,
                "report_scope": "complete",
            },
        )

        self.assertRedirects(response, reverse("clinic:patient_detail", args=[self.patient.pk]))
        report_types = set(GeneratedReport.objects.filter(encounter=self.encounter).values_list("report_type", flat=True))
        self.assertIn(ReportType.COMPLETO, report_types)
        self.assertIn(ReportType.MUTUAL, report_types)


class CalendarEditingTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(username="calendar-edit", password="secret123")
        grant_clinic_permissions(self.user, "manage_agenda")
        self.client.force_login(self.user)
        self.patient = Patient.objects.create(full_name="ARRIETA, LIDIA", dni="17124122")
        self.encounter = Encounter.objects.create(
            patient=self.patient,
            encounter_date=date(2026, 6, 12),
            encounter_time=time(10, 0),
            study_type=StudyType.CICLOMETRIA,
            coverage_type=CoverageType.PARTICULAR,
            status=EncounterStatus.PENDIENTE,
            created_by=self.user,
            updated_by=self.user,
        )

    def test_calendar_allows_updating_coverage(self):
        response = self.client.post(
            reverse("clinic:calendar"),
            {
                "action": "update_encounter_field",
                "encounter_id": self.encounter.pk,
                "field_name": "coverage_type",
                "value": "Mutual",
                "month": "2026-06",
                "date": "2026-06-12",
            },
        )

        self.assertRedirects(response, f"{reverse('clinic:calendar')}?month=2026-06&date=2026-06-12")
        self.encounter.refresh_from_db()
        self.assertEqual(self.encounter.coverage_type, CoverageType.MUTUAL)

    def test_calendar_detail_renders_inline_edit_selects(self):
        response = self.client.get(f"{reverse('clinic:calendar')}?month=2026-06&date=2026-06-12")

        self.assertEqual(response.status_code, 200)
        html = response.content.decode()
        self.assertIn('name="field_name" value="coverage_type"', html)
        self.assertIn('name="field_name" value="study_type"', html)
        self.assertIn("Editar ficha", html)


class StatisticsMonthNavigationTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(username="stats-months", password="secret123")
        grant_clinic_permissions(self.user, "view_clinical_statistics")
        self.client.force_login(self.user)
        patient = Patient.objects.create(full_name="MES, PRUEBA", dni="11111111")
        Encounter.objects.create(
            patient=patient,
            encounter_date=date(2026, 5, 20),
            encounter_time=time(9, 0),
            study_type=StudyType.CICLOMETRIA,
            coverage_type=CoverageType.MUTUAL,
            status=EncounterStatus.PENDIENTE,
            created_by=self.user,
            updated_by=self.user,
        )
        Encounter.objects.create(
            patient=patient,
            encounter_date=date(2026, 6, 10),
            encounter_time=time(10, 0),
            study_type=StudyType.CICLOMETRIA,
            coverage_type=CoverageType.PARTICULAR,
            status=EncounterStatus.PENDIENTE,
            created_by=self.user,
            updated_by=self.user,
        )

    def test_statistics_allows_browsing_previous_month(self):
        response = self.client.get(f"{reverse('clinic:statistics')}?month=2026-05")

        self.assertEqual(response.status_code, 200)
        html = response.content.decode()
        self.assertIn("mayo 2026", html.lower())
        self.assertIn("Mes elegido", html)

    def test_statistics_caps_future_month_to_current(self):
        response = self.client.get(f"{reverse('clinic:statistics')}?month=2099-01")

        self.assertEqual(response.status_code, 200)
        html = response.content.decode()
        self.assertIn("Mes actual", html)


class ClinicalAccessControlTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(username="read-only", password="secret123")
        self.client.force_login(self.user)
        self.patient = Patient.objects.create(full_name="SOLO, LECTURA", dni="55666777")
        self.encounter = Encounter.objects.create(
            patient=self.patient,
            encounter_date=date(2026, 7, 13),
            study_type=StudyType.CICLOMETRIA,
            coverage_type=CoverageType.PARTICULAR,
            created_by=self.user,
            updated_by=self.user,
        )

    def test_read_only_user_cannot_mutate_dashboard(self):
        response = self.client.post(
            reverse("clinic:dashboard"),
            {
                "action": "inline_update",
                "encounter_id": self.encounter.pk,
                "field_name": "patient_name",
                "value": "ALTERADO",
            },
        )

        self.assertEqual(response.status_code, 403)
        self.patient.refresh_from_db()
        self.assertEqual(self.patient.full_name, "SOLO, LECTURA")

    def test_read_only_user_cannot_delete_patient(self):
        response = self.client.post(reverse("clinic:patient_delete", args=[self.patient.pk]))

        self.assertEqual(response.status_code, 403)
        self.assertTrue(Patient.objects.filter(pk=self.patient.pk).exists())

    def test_specialized_sections_require_their_permissions(self):
        self.assertEqual(self.client.get(reverse("clinic:doctor_review_list")).status_code, 403)
        self.assertEqual(self.client.get(reverse("clinic:statistics")).status_code, 403)
        self.assertEqual(self.client.get(reverse("clinic:recycle_bin")).status_code, 403)
