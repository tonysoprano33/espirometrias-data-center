from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .models import AttachmentKind


DEFAULT_DOCTOR = "DR. GUSTAVO PIGUILLEM"


def sanitizar_nombre_archivo(texto: str) -> str:
    limpio = re.sub(r'[<>:"/\\|?*]+', " ", str(texto))
    limpio = re.sub(r"\s+", " ", limpio).strip()
    return (limpio or "SIN_NOMBRE").replace(" ", "_")


def limpiar_entero(valor, default: str = "0") -> str:
    digits = re.sub(r"\D", "", str(valor))
    return digits or default


def formatear_dni(dni_raw) -> str:
    dni_raw = re.sub(r"\D", "", str(dni_raw))
    if len(dni_raw) == 0:
        return "0"
    if len(dni_raw) <= 3:
        return dni_raw
    if len(dni_raw) <= 6:
        return f"{dni_raw[:-3]}.{dni_raw[-3:]}"
    return f"{dni_raw[:-6]}.{dni_raw[-6:-3]}.{dni_raw[-3:]}"


def normalizar_medico(nombre) -> str:
    limpio = re.sub(r"\s+", " ", str(nombre or "").strip().upper())
    if not limpio:
        return DEFAULT_DOCTOR
    limpio = re.sub(r"^DR\.?\s*", "", limpio)
    return f"DR. {limpio}"


def normalizar_patron(valor) -> str:
    texto = str(valor or "").strip().lower()
    mapping = {
        "normal": "Normal",
        "obstructivo": "Obstructivo",
        "restrictivo": "Restrictivo",
        "mixto": "Mixto",
    }
    return mapping.get(texto, "Normal")


def construir_informe_espirometria(patron: str, grado_obst: str, grado_rest: str) -> str:
    if patron == "Normal":
        return "El paciente presenta resultados normales."
    if patron == "Obstructivo":
        return f"El paciente presenta déficit respiratorio (obstrucción {grado_obst}) a las pequeñas vías respiratorias aéreas."
    if patron == "Restrictivo":
        return f"El paciente presenta déficit respiratorio (restricción {grado_rest}) a las vías respiratorias aéreas."
    grado_rest_text = str(grado_rest or "").strip().capitalize()
    grado_obst_text = str(grado_obst or "").strip().capitalize()
    return (
        "El paciente presenta déficit respiratorio con patrón mixto:\n\n"
        f"Restricción {grado_rest_text}.\n\n"
        f"Obstrucción {grado_obst_text} a las vías respiratorias aéreas."
    )


def agregar_borde_parrafo(paragraph):
    p = paragraph._element
    p_pr = p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "12")
        border.set(qn("w:space"), "8")
        border.set(qn("w:color"), "4A90E2")
        p_bdr.append(border)
    p_pr.append(p_bdr)


def interpolar_valores(minimo, maximo, cantidad: int = 7) -> list[int]:
    minimo = int(minimo or 0)
    maximo = int(maximo or 0)
    if cantidad <= 2:
        return [minimo, maximo]
    paso = (maximo - minimo) / (cantidad - 1)
    return [round(minimo + (paso * i)) for i in range(cantidad)]


def build_walk_test_assessment(
    so2_rest=None,
    so2_post=None,
    *,
    completed: bool = True,
    stopped: bool = False,
    symptoms: bool = False,
):
    def parse_optional_int(value):
        try:
            if value in (None, ""):
                return None
            return int(value)
        except (TypeError, ValueError):
            return None

    rest = parse_optional_int(so2_rest)
    post = parse_optional_int(so2_post)
    reasons = []
    drop = None if rest is None or post is None else rest - post

    if post is not None and post <= 88:
        reasons.append("desaturacion al esfuerzo")
    if drop is not None and drop >= 4:
        reasons.append("caida significativa de SO2")
    if not completed:
        reasons.append("marcha no completada")
    if stopped:
        reasons.append("se detuvo durante la marcha")
    if symptoms:
        reasons.append("presento sintomas")

    if reasons:
        return {
            "is_normal": False,
            "label": "PRUEBA NO NORMAL",
            "detail": "",
            "tone": "alert",
        }

    return {
        "is_normal": True,
        "label": "PRUEBA NORMAL",
        "detail": "",
        "tone": "ok",
    }


def crear_encabezado(doc: Document):
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    header_box = doc.add_paragraph()
    header_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    agregar_borde_parrafo(header_box)
    header_box.paragraph_format.space_before = Pt(6)
    header_box.paragraph_format.space_after = Pt(6)
    run1 = header_box.add_run("CENTRO RESPIRATORIO INTEGRAL\n")
    run1.bold = True
    run1.font.name = "Times New Roman"
    run1.font.size = Pt(14)
    run1.font.color.rgb = RGBColor(40, 60, 90)
    run2 = header_box.add_run("MARCONI 147 - TEL: 02657-705270\n")
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(30, 30, 30)
    run3 = header_box.add_run("VILLA MERCEDES (SAN LUIS)")
    run3.font.name = "Times New Roman"
    run3.font.size = Pt(11)
    run3.font.color.rgb = RGBColor(30, 30, 30)


def agregar_fecha(doc: Document, fecha_texto: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(fecha_texto)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def agregar_datos_paciente(doc: Document, nombre: str, dni: str, deriva: str):
    for label_text, value in [("PACIENTE: ", nombre), ("DNI: ", dni), ("DERIVA: ", deriva)]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        label = p.add_run(label_text)
        label.bold = True
        label.underline = True
        label.font.name = "Times New Roman"
        label.font.size = Pt(14)
        label.font.color.rgb = RGBColor(40, 60, 90)
        value_run = p.add_run(value)
        value_run.font.name = "Times New Roman"
        if label_text == "DNI: ":
            value_run.font.size = Pt(14)
            value_run.bold = True
        elif label_text == "PACIENTE: ":
            value_run.font.size = Pt(16)
            value_run.bold = True
        else:
            value_run.font.size = Pt(14)


def agregar_firma(doc: Document, as_footer: bool = False):
    signature_text = "DR. PIGUILLEM GUSTAVO GABRIEL\nMAT. 2083\nESP. EN VÍAS RESPIRATORIAS"
    if as_footer:
        footer = doc.sections[-1].footer
        for p in footer.paragraphs:
            p.clear()
        p = footer.add_paragraph()
    else:
        p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(signature_text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(9)


def agregar_salto_pagina(doc: Document):
    doc.add_page_break()


def agregar_seccion_espirometria(doc: Document, so2: str, fc: str, informe: str, es_normal: bool, broncodilatador_positivo: bool = False):
    t = doc.add_paragraph()
    t.paragraph_format.space_before = Pt(18)
    t.paragraph_format.space_after = Pt(12)
    run = t.add_run("Resultado Espirometría Computarizada:")
    run.bold = True
    run.underline = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(40, 60, 90)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(informe)
    r.font.name = "Times New Roman"
    r.font.size = Pt(16)
    r.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5

    if broncodilatador_positivo:
        doc.add_paragraph()
        bronco_p = doc.add_paragraph()
        bronco_p.paragraph_format.space_before = Pt(12)
        bronco_p.paragraph_format.space_after = Pt(12)
        bronco_run = bronco_p.add_run("Test de Broncodilatador: POSITIVO")
        bronco_run.font.name = "Times New Roman"
        bronco_run.font.size = Pt(15)
        bronco_run.bold = True
        bronco_run.font.color.rgb = RGBColor(200, 50, 50)

    val = doc.add_paragraph()
    val.paragraph_format.space_before = Pt(24)
    val.paragraph_format.space_after = Pt(12)
    val_run = val.add_run(f"SO2: {so2}%          FC: {fc}%")
    val_run.bold = True
    val_run.font.name = "Times New Roman"
    val_run.font.size = Pt(15)

    if not es_normal:
        rec = doc.add_paragraph()
        rec.paragraph_format.space_before = Pt(20)
        rec.paragraph_format.space_after = Pt(10)
        rec_run = rec.add_run("Por antecedentes clínicos del paciente, sugiero control.")
        rec_run.italic = True
        rec_run.font.name = "Times New Roman"
        rec_run.font.size = Pt(14)
        rec_run.bold = True


def agregar_seccion_caminata(
    doc: Document,
    so2_vals: list[int],
    fc_vals: list[int],
    distancia: str,
    concluida: bool,
    detuvo: bool,
    sintomas: bool,
    borg_vals: list[int],
    walk_label: str,
    walk_detail: str = "",
):
    titulo = doc.add_paragraph()
    run = titulo.add_run("PRUEBA DE LOS 6 Y 12 MINUTOS")
    run.bold = True
    run.underline = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    desc = doc.add_paragraph()
    desc_run = desc.add_run(
        "Se realizó test de la marcha con monitoreo continuo; en reposo, durante la marcha (6 min.) "
        "y en la recuperación (2 min.). Los parámetros registrados fueron: Saturación de O2, frecuencia de pulso, "
        "esfuerzo percibido por medio de escala de Borg y la distancia recorrida a la finalización."
    )
    desc_run.font.name = "Times New Roman"
    desc_run.font.size = Pt(10)
    desc.paragraph_format.space_after = Pt(6)
    desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    com = doc.add_paragraph("COMENTARIOS:")
    if com.runs:
        com.runs[0].bold = True
        com.runs[0].underline = True
        com.runs[0].font.name = "Times New Roman"
        com.runs[0].font.size = Pt(10)

    preguntas = [
        f"Distancia recorrida: {distancia} mts.",
        f"Realizó correctamente la marcha: {'Sí' if concluida else 'No'}.",
        f"Se detuvo durante la marcha: {'Sí' if detuvo else 'No'}.",
        f"Presentó algún síntoma al final de la marcha: {'Sí' if sintomas else 'No'}.",
    ]
    for q in preguntas:
        p = doc.add_paragraph(q)
        p.paragraph_format.left_indent = Inches(0.3)
        if p.runs:
            p.runs[0].font.name = "Times New Roman"
            p.runs[0].font.size = Pt(10)

    tabla = doc.add_table(rows=8, cols=4)
    tabla.style = "Table Grid"
    hdr = ["MINUTOS", "SO2", "FC", "ESC.BORG"]
    for i, h in enumerate(hdr):
        cell = tabla.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i in range(7):
        tabla.rows[i + 1].cells[0].text = str(i)
        tabla.rows[i + 1].cells[1].text = str(so2_vals[i])
        tabla.rows[i + 1].cells[2].text = str(fc_vals[i])
        tabla.rows[i + 1].cells[3].text = str(borg_vals[i])
        for j in range(4):
            for para in tabla.rows[i + 1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    res = doc.add_paragraph()
    run = res.add_run("PRUEBA DE LOS 6 Y 12 MINUTOS: ")
    run.bold = True
    run.underline = True
    run.font.name = "Times New Roman"
    doc.add_paragraph()
    run.font.size = Pt(11)
    normal_run = res.add_run(walk_label)
    normal_run.font.name = "Times New Roman"
    normal_run.font.size = Pt(11)


def crear_informe_mutual(
    nombre: str,
    dni: str,
    fecha: str,
    deriva: str,
    so2: str,
    fc: str,
    so2_vals: list[int],
    fc_vals: list[int],
    borg_vals: list[int],
    distancia: str,
    concluida: bool,
    detuvo: bool,
    sintomas: bool,
    informe_espiro: str,
    patron: str,
    grado_obst: str,
    grado_rest: str,
    walk_label: str,
    walk_detail: str = "",
) -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    header_box = doc.add_paragraph()
    header_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    agregar_borde_parrafo(header_box)
    header_box.paragraph_format.space_before = Pt(4)
    header_box.paragraph_format.space_after = Pt(4)
    run1 = header_box.add_run("CENTRO RESPIRATORIO INTEGRAL\n")
    run1.bold = True
    run1.font.size = Pt(13)
    run1.font.color.rgb = RGBColor(40, 60, 90)
    run2 = header_box.add_run("MARCONI 147 - TEL: 02657-705270\n")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(30, 30, 30)
    run3 = header_box.add_run("VILLA MERCEDES (SAN LUIS)")
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(30, 30, 30)

    p_fecha = doc.add_paragraph()
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_fecha.paragraph_format.space_before = Pt(4)
    p_fecha.paragraph_format.space_after = Pt(6)
    run_fecha = p_fecha.add_run(fecha)
    run_fecha.bold = True
    run_fecha.font.size = Pt(11)

    for label_text, value in [("PACIENTE: ", nombre), ("DNI: ", dni), ("DERIVA: ", deriva)]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        label = p.add_run(label_text)
        label.bold = True
        label.underline = True
        label.font.size = Pt(11)
        label.font.color.rgb = RGBColor(40, 60, 90)
        p.add_run(value).font.size = Pt(11)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(8)

    titulo = doc.add_paragraph()
    titulo.paragraph_format.space_before = Pt(6)
    titulo.paragraph_format.space_after = Pt(4)
    run_titulo = titulo.add_run("PRUEBA DE LOS 6 Y 12 MINUTOS")
    run_titulo.bold = True
    run_titulo.underline = True
    run_titulo.font.size = Pt(12)

    desc = doc.add_paragraph()
    desc.paragraph_format.space_before = Pt(2)
    desc.paragraph_format.space_after = Pt(6)
    desc_run = desc.add_run(
        "Se realizó test de la marcha con monitoreo continuo; en reposo, durante la marcha (6 min.) "
        "y en la recuperación (2 min.). Los parámetros registrados fueron: Saturación de O2, frecuencia de pulso, "
        "esfuerzo percibido por medio de escala de Borg y la distancia recorrida a la finalización."
    )
    desc_run.font.size = Pt(10)
    desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    com = doc.add_paragraph("COMENTARIOS:")
    com.paragraph_format.space_before = Pt(4)
    com.paragraph_format.space_after = Pt(2)
    if com.runs:
        com.runs[0].bold = True
        com.runs[0].underline = True
        com.runs[0].font.size = Pt(10)

    preguntas = [
        f"Distancia recorrida: {distancia} mts.",
        f"Realizó correctamente la marcha: {'Sí' if concluida else 'No'}.",
        f"Se detuvo durante la marcha: {'Sí' if detuvo else 'No'}.",
        f"Presentó algún síntoma al final de la marcha: {'Sí' if sintomas else 'No'}.",
    ]
    for q in preguntas:
        p = doc.add_paragraph(q)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.25)
        if p.runs:
            p.runs[0].font.size = Pt(10)

    tabla = doc.add_table(rows=8, cols=4)
    tabla.style = "Table Grid"
    hdr = ["MINUTOS", "SO2", "FC", "ESC.BORG"]
    for i, h in enumerate(hdr):
        cell = tabla.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i in range(7):
        tabla.rows[i + 1].cells[0].text = str(i)
        tabla.rows[i + 1].cells[1].text = str(so2_vals[i])
        tabla.rows[i + 1].cells[2].text = str(fc_vals[i])
        tabla.rows[i + 1].cells[3].text = str(borg_vals[i])
        for j in range(4):
            for para in tabla.rows[i + 1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_so2_fc = doc.add_paragraph()
    p_so2_fc.paragraph_format.space_before = Pt(6)
    p_so2_fc.paragraph_format.space_after = Pt(6)
    run_so2_fc = p_so2_fc.add_run(f"SO2: {so2}%    FC: {fc}")
    run_so2_fc.bold = True
    run_so2_fc.font.size = Pt(11)

    res = doc.add_paragraph()
    res.paragraph_format.space_before = Pt(4)
    res.paragraph_format.space_after = Pt(6)
    run = res.add_run("PRUEBA DE LOS 6 Y 12 MINUTOS: ")
    run.bold = True
    run.underline = True
    run.font.size = Pt(11)
    res.add_run(walk_label).font.size = Pt(11)

    esp_titulo = doc.add_paragraph()
    esp_titulo.paragraph_format.space_before = Pt(6)
    esp_titulo.paragraph_format.space_after = Pt(4)
    esp_run = esp_titulo.add_run("Resultado Espirometría Computarizada:")
    esp_run.bold = True
    esp_run.underline = True
    esp_run.font.size = Pt(11)
    esp_run.font.color.rgb = RGBColor(40, 60, 90)

    esp_result = doc.add_paragraph()
    esp_result.paragraph_format.space_before = Pt(2)
    esp_result.paragraph_format.space_after = Pt(6)
    esp_result_run = esp_result.add_run(informe_espiro)
    esp_result_run.font.size = Pt(11)

    cvl_titulo = doc.add_paragraph()
    cvl_titulo.paragraph_format.space_before = Pt(4)
    cvl_titulo.paragraph_format.space_after = Pt(2)
    run_cvl = cvl_titulo.add_run("Capacidad Vital Lenta:")
    run_cvl.bold = True
    run_cvl.underline = True
    run_cvl.font.size = Pt(11)
    run_cvl.font.color.rgb = RGBColor(40, 60, 90)

    grado_obst = (grado_obst or "").strip().lower()
    grado_rest = (grado_rest or "").strip().lower()
    if patron == "Normal":
        resultado_cvl = "Normal"
    elif patron == "Obstructivo":
        if grado_obst == "leve":
            resultado_cvl = "Levemente disminuida"
        elif grado_obst in {"moderado", "moderada"}:
            resultado_cvl = "Moderadamente disminuida"
        elif grado_obst == "moderadamente severa":
            resultado_cvl = "Moderadamente a severamente disminuida"
        else:
            resultado_cvl = "Severamente disminuida"
    elif patron == "Restrictivo":
        if grado_rest == "leve":
            resultado_cvl = "Levemente reducida"
        elif grado_rest in {"moderado", "moderada"}:
            resultado_cvl = "Moderadamente reducida"
        elif grado_rest == "moderadamente severa":
            resultado_cvl = "Moderadamente a severamente reducida"
        else:
            resultado_cvl = "Severamente reducida"
    else:
        resultado_cvl = "Reducida (patrón mixto)"

    p_cvl = doc.add_paragraph()
    p_cvl.paragraph_format.space_before = Pt(2)
    p_cvl.paragraph_format.space_after = Pt(4)
    r = p_cvl.add_run(resultado_cvl)
    r.font.size = Pt(11)
    r.bold = True

    agregar_firma(doc, as_footer=True)
    return doc


@dataclass
class GeneratedArtifact:
    report_type: str
    filename: str
    bytes_content: bytes
    file_kind: str = AttachmentKind.INFORME_DOCX
    mime_type: str = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _document_to_bytes(doc: Document) -> bytes:
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _latest_pdf_attachment(encounter):
    return (
        encounter.attachments.filter(file_kind__in=[AttachmentKind.PDF_RESULTADO, AttachmentKind.FOTO_RESULTADO])
        .order_by("-created_at")
        .first()
    )


def _pdf_artifact_from_original(encounter, nombre_archivo_seguro: str, fecha_archivo: str) -> GeneratedArtifact:
    attachment = _latest_pdf_attachment(encounter)
    if not attachment or not getattr(attachment, "file", None):
        raise ValueError("Para imprimir una espirometria sola, primero subi el PDF o una foto original del equipo.")

    with open(attachment.file.path, "rb") as pdf_file:
        bytes_content = pdf_file.read()

    if attachment.file_kind == AttachmentKind.FOTO_RESULTADO:
        mime_type = str(getattr(attachment, "mime_type", "") or "image/jpeg")
        extension = ".jpg"
        original_name = str(getattr(attachment, "original_name", "") or "")
        if "." in original_name:
            extension = "." + original_name.split(".")[-1].lower()
        return GeneratedArtifact(
            report_type="Espirometria",
            filename=f"Espirometria_{nombre_archivo_seguro}_{fecha_archivo}{extension}",
            bytes_content=bytes_content,
            file_kind=AttachmentKind.OTRO,
            mime_type=mime_type,
        )

    return GeneratedArtifact(
        report_type="Espirometria",
        filename=f"Espirometria_{nombre_archivo_seguro}_{fecha_archivo}.pdf",
        bytes_content=bytes_content,
        file_kind=AttachmentKind.INFORME_PDF,
        mime_type="application/pdf",
    )


def _single_page_spirometry_artifact(
    *,
    nombre: str,
    dni: str,
    deriva: str,
    fecha_impresion: str,
    so2: str,
    fc: str,
    informe: str,
    es_normal: bool,
    broncodilatador_positivo: bool,
    nombre_archivo_seguro: str,
    fecha_archivo: str,
) -> GeneratedArtifact:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    crear_encabezado(doc)
    agregar_fecha(doc, fecha_impresion)
    agregar_datos_paciente(doc, nombre, dni, deriva)
    agregar_seccion_espirometria(doc, so2, fc, informe, es_normal, broncodilatador_positivo)
    agregar_firma(doc, as_footer=True)

    return GeneratedArtifact(
        report_type="Espirometria",
        filename=f"Informe_Espirometria_{nombre_archivo_seguro}_{fecha_archivo}.docx",
        bytes_content=_document_to_bytes(doc),
    )


def _encounter_has_walk_data(encounter) -> bool:
    if getattr(encounter, "study_type", "") != "Ciclometria":
        return False
    walk = getattr(encounter, "walk_test", None)
    vitals = getattr(encounter, "vital_signs", None)
    if not walk and not vitals:
        return False
    return any(
        [
            getattr(walk, "distance_meters", None),
            getattr(walk, "completed", False),
            getattr(walk, "stopped", False),
            getattr(walk, "symptoms", False),
            getattr(walk, "borg_final", None) not in (None, 0),
            getattr(vitals, "so2_post", None),
            getattr(vitals, "fc_post", None),
        ]
    )


def build_reports_for_encounter(encounter) -> list[GeneratedArtifact]:
    patient = encounter.patient
    vital = getattr(encounter, "vital_signs", None)
    walk = getattr(encounter, "walk_test", None)
    result = getattr(encounter, "spirometry_result", None)

    nombre = str(patient.full_name or "").strip().upper()
    if not nombre:
        raise ValueError("La atención no tiene nombre de paciente.")

    fecha_impresion = encounter.encounter_date.strftime("%d/%m/%Y")
    dni = formatear_dni(patient.dni)
    deriva = normalizar_medico(getattr(encounter.referring_physician, "full_name", DEFAULT_DOCTOR))
    so2 = limpiar_entero(getattr(vital, "so2_rest", "0"))
    fc = limpiar_entero(getattr(vital, "fc_rest", "0"))
    patron = normalizar_patron(getattr(result, "respiratory_pattern", "Normal"))
    es_normal = patron == "Normal"
    grado_obst = (getattr(result, "obstruction_grade", "") or "Leve").strip().lower()
    grado_rest = (getattr(result, "restriction_grade", "") or "Leve").strip().lower()
    broncodilatador_positivo = bool(getattr(result, "bronchodilator_positive", False))
    informe = construir_informe_espirometria(patron, grado_obst, grado_rest)
    incluir_caminata = _encounter_has_walk_data(encounter)

    so2_vals = []
    fc_vals = []
    borg_vals = []
    distancia = str(getattr(walk, "distance_meters", "") or "200")
    concluida = bool(getattr(walk, "completed", True))
    detuvo = bool(getattr(walk, "stopped", False))
    sintomas = bool(getattr(walk, "symptoms", False))
    walk_assessment = build_walk_test_assessment(
        getattr(vital, "so2_rest", None),
        getattr(vital, "so2_post", None),
        completed=concluida,
        stopped=detuvo,
        symptoms=sintomas,
    )

    nombre_archivo_seguro = sanitizar_nombre_archivo(nombre)
    fecha_archivo = fecha_impresion.replace("/", "-")
    artifacts = []

    if not incluir_caminata:
        artifacts.append(
            _single_page_spirometry_artifact(
                nombre=nombre,
                dni=dni,
                deriva=deriva,
                fecha_impresion=fecha_impresion,
                so2=so2,
                fc=fc,
                informe=informe,
                es_normal=es_normal,
                broncodilatador_positivo=broncodilatador_positivo,
                nombre_archivo_seguro=nombre_archivo_seguro,
                fecha_archivo=fecha_archivo,
            )
        )
        return artifacts

    doc_normal = Document()
    for section in doc_normal.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    crear_encabezado(doc_normal)
    agregar_fecha(doc_normal, fecha_impresion)
    agregar_datos_paciente(doc_normal, nombre, dni, deriva)
    agregar_seccion_espirometria(doc_normal, so2, fc, informe, es_normal, broncodilatador_positivo)

    so2_reposo = int(so2)
    so2_regreso = int(limpiar_entero(getattr(vital, "so2_post", "100"), "100"))
    fc_reposo = int(fc)
    fc_maximo = int(limpiar_entero(getattr(vital, "fc_post", "120"), "120"))
    borg_final = int(getattr(walk, "borg_final", 0) or 0)

    so2_vals = interpolar_valores(so2_reposo, so2_regreso, 7)
    fc_vals = interpolar_valores(fc_reposo, fc_maximo, 7)
    # Match the manual .exe path in E:\espiro\main.py (App.crear_documento_logica):
    # when Borg final is 0, the table still ends with 1 at minute 6.
    borg_vals = interpolar_valores(0, borg_final, 7) if borg_final > 0 else [0, 0, 0, 0, 0, 0, 1]

    agregar_salto_pagina(doc_normal)
    crear_encabezado(doc_normal)
    agregar_fecha(doc_normal, fecha_impresion)
    agregar_datos_paciente(doc_normal, nombre, dni, deriva)
    agregar_seccion_caminata(
        doc_normal,
        so2_vals,
        fc_vals,
        distancia,
        concluida,
        detuvo,
        sintomas,
        borg_vals,
        walk_assessment["label"],
        walk_assessment["detail"],
    )
    agregar_firma(doc_normal, as_footer=True)

    artifacts.append(
        GeneratedArtifact(
            report_type="Completo",
            filename=f"Informe_Completo_{nombre_archivo_seguro}_{fecha_archivo}.docx",
            bytes_content=_document_to_bytes(doc_normal),
        )
    )

    incluir_mutual = encounter.coverage_type == "Mutual"
    if incluir_mutual and incluir_caminata:
        doc_mutual = crear_informe_mutual(
            nombre,
            dni,
            fecha_impresion,
            deriva,
            so2,
            fc,
            so2_vals,
            fc_vals,
            borg_vals,
            distancia,
            concluida,
            detuvo,
            sintomas,
            informe,
            patron,
            grado_obst,
            grado_rest,
            walk_assessment["label"],
            walk_assessment["detail"],
        )
        artifacts.append(
            GeneratedArtifact(
                report_type="Mutual",
                filename=f"MUTUAL_Informe_{nombre_archivo_seguro}_{fecha_archivo}.docx",
                bytes_content=_document_to_bytes(doc_mutual),
            )
        )

    return artifacts

